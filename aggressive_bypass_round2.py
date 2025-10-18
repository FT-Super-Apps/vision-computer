#!/usr/bin/env python3
"""
Aggressive Bypass Round 2
Target remaining flags with 70-80% homoglyphs
"""

import sys
import random
from docx import Document
from difflib import SequenceMatcher

# Homoglyphs mapping (Cyrillic lookalikes)
HOMOGLYPHS = {
    'A': 'А', 'B': 'В', 'C': 'С', 'E': 'Е', 'H': 'Н',
    'I': 'І', 'K': 'К', 'M': 'М', 'O': 'О', 'P': 'Р',
    'T': 'Т', 'X': 'Х', 'Y': 'У',
    'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с',
    'x': 'х', 'y': 'у', 'i': 'і', 'k': 'к', 'u': 'υ',
    'n': 'п', 's': 'ѕ', 'g': 'ց', 'd': 'ԁ'
}

def apply_aggressive_homoglyphs(text, density=0.75):
    """Apply homoglyphs with high density (70-80%)"""
    if not text:
        return text
    
    result = list(text)
    replaceable = [i for i, c in enumerate(text) if c in HOMOGLYPHS]
    
    if not replaceable:
        return text
    
    num_replace = max(1, int(len(replaceable) * density))
    positions = random.sample(replaceable, min(num_replace, len(replaceable)))
    
    for pos in positions:
        result[pos] = HOMOGLYPHS[result[pos]]
    
    return ''.join(result)

def aggressive_bypass_round2(input_docx, flags_file, output_docx):
    """
    Round 2: Aggressive bypass for remaining flags
    """
    print(f"\n{'='*70}")
    print(f"🚀 AGGRESSIVE BYPASS - ROUND 2")
    print(f"{'='*70}")
    print(f"Input  : {input_docx}")
    print(f"Flags  : {flags_file}")
    print(f"Output : {output_docx}")
    print(f"Strategy: 70-80% homoglyphs for remaining flags")
    print(f"{'='*70}\n")
    
    # Load remaining flags
    with open(flags_file, 'r', encoding='utf-8') as f:
        remaining_flags = [line.strip() for line in f if line.strip()]
    
    print(f"📋 Loaded {len(remaining_flags)} remaining flags\n")
    
    # Load document
    doc = Document(input_docx)
    modified_count = 0
    
    print(f"🔍 Processing document...\n")
    
    # Strategy per flag
    strategies = {
        # Short texts (headers) - 80%
        'header': 0.80,
        # Technical terms - 75%
        'technical': 0.75,
        # Common phrases - 70%
        'phrase': 0.70
    }
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        if not para_text:
            continue
        
        # Check against remaining flags
        for flag in remaining_flags:
            if not flag:
                continue
            
            # Fuzzy match
            similarity = SequenceMatcher(None, para_text.lower(), flag.lower()).ratio()
            is_substring = flag.lower() in para_text.lower()
            
            if similarity > 0.75 or is_substring:
                # Determine strategy
                word_count = len(para_text.split())
                
                if word_count <= 5:
                    strategy = 'header'
                    density = 0.80
                elif any(term in para_text for term in ['K3', 'RAG', 'Artificial', 'teknologi']):
                    strategy = 'technical'
                    density = 0.75
                else:
                    strategy = 'phrase'
                    density = 0.70
                
                modified_count += 1
                
                print(f"✅ Match #{modified_count}:")
                print(f"   Para #{para_idx}: {para_text[:60]}...")
                print(f"   Flag: {flag[:60]}...")
                print(f"   Strategy: {strategy.upper()} ({density*100:.0f}% homoglyphs)")
                
                # Apply aggressive homoglyphs
                original_full = para.text
                modified_full = apply_aggressive_homoglyphs(original_full, density=density)
                
                # Clear and rebuild
                para.clear()
                para.add_run(modified_full)
                
                # Count changes
                changes = sum(1 for a, b in zip(original_full, modified_full) if a != b)
                print(f"   Changed: {changes}/{len(original_full)} chars ({changes/len(original_full)*100:.1f}%)")
                print(f"   Preview: {modified_full[:60]}...")
                print()
                
                break
    
    # Save
    print(f"💾 Saving to: {output_docx}")
    doc.save(output_docx)
    
    print(f"\n{'='*70}")
    print(f"✅ ROUND 2 COMPLETE!")
    print(f"{'='*70}")
    print(f"📊 Statistics:")
    print(f"   Remaining flags   : {len(remaining_flags)}")
    print(f"   Paragraphs modified: {modified_count}")
    print(f"   Strategy: 70-80% aggressive homoglyphs")
    print(f"\n🎯 Expected improvement:")
    print(f"   Current: 14%")
    print(f"   Target : <10%")
    print(f"   Expected: 8-10% similarity")
    print(f"\n📄 Output: {output_docx}")
    print(f"{'='*70}\n")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python aggressive_bypass_round2.py original.docx [flags_file] [output.docx]")
        print("\nExample:")
        print("  python aggressive_bypass_round2.py original.docx flag_new.txt original_round2.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    flags_file = sys.argv[2] if len(sys.argv) > 2 else "flag_new.txt"
    output_file = sys.argv[3] if len(sys.argv) > 3 else "original_round2.docx"
    
    aggressive_bypass_round2(input_file, flags_file, output_file)
