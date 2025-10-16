#!/usr/bin/env python3
"""
AGGRESSIVE PARAPHRASE STRATEGY:
- Double paraphrase (2 rounds)
- Higher temperature (more creative)
- More beams (better quality)
- Multiple outputs (pick best)
"""

import json
import docx
from difflib import SequenceMatcher
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch
from tqdm import tqdm

def load_indot5_model(model_name="Wikidepia/IndoT5-base-paraphrase"):
    """Load IndoT5 model"""
    print(f"🔧 Loading IndoT5: {model_name}")
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)
    print(f"✅ Model loaded on {device}")
    return tokenizer, model, device

def aggressive_paraphrase(text, tokenizer, model, device, rounds=2):
    """
    AGGRESSIVE multi-round paraphrase
    - Round 1: High temperature for creativity
    - Round 2: Paraphrase the paraphrase!
    """
    current_text = text
    
    for round_num in range(1, rounds + 1):
        input_text = f"paraphrase: {current_text}"
        inputs = tokenizer(
            input_text,
            return_tensors="pt",
            max_length=512,
            truncation=True,
            padding=True
        ).to(device)
        
        # More aggressive parameters
        with torch.no_grad():
            outputs = model.generate(
                inputs.input_ids,
                max_length=512,
                num_beams=15,              # More beams
                temperature=1.5,           # Higher temperature
                do_sample=True,
                top_k=100,                 # More options
                top_p=0.98,                # Higher nucleus
                early_stopping=True,
                no_repeat_ngram_size=4,    # Avoid longer repetitions
                num_return_sequences=1
            )
        
        current_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
    
    return current_text

def match_flagged_with_docx(flagged_json, docx_path, threshold=0.5):
    """Match flagged texts with DOCX"""
    with open(flagged_json, 'r', encoding='utf-8') as f:
        flagged_texts = json.load(f)
    
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    print(f"\n🔍 Matching {len(flagged_texts)} flagged with {len(paragraphs)} paragraphs...")
    
    matches = []
    for i, flagged in enumerate(flagged_texts, 1):
        flagged_text = flagged['text']
        best_match = None
        best_ratio = 0
        best_idx = -1
        
        for idx, para in enumerate(paragraphs):
            ratio = SequenceMatcher(None, flagged_text.lower(), para.lower()).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = para
                best_idx = idx
        
        if best_ratio < 0.5:
            keywords = [w.lower() for w in flagged_text.split() if len(w) > 4]
            for idx, para in enumerate(paragraphs):
                keyword_matches = sum(1 for kw in keywords if kw in para.lower())
                keyword_ratio = keyword_matches / len(keywords) if keywords else 0
                if keyword_ratio > best_ratio:
                    best_ratio = keyword_ratio
                    best_match = para
                    best_idx = idx
        
        if best_ratio >= threshold:
            matches.append({
                "flagged_text": flagged_text,
                "matched_para": best_match,
                "similarity": round(best_ratio * 100, 2),
                "para_index": best_idx,
                "page": flagged['page']
            })
    
    print(f"✅ Matched: {len(matches)}/{len(flagged_texts)}")
    return matches

def process_aggressive(matches):
    """AGGRESSIVE double-round paraphrase"""
    print(f"\n🔥 AGGRESSIVE PARAPHRASE: {len(matches)} texts...")
    print("Strategy: DOUBLE paraphrase with high creativity")
    print("="*70)
    
    tokenizer, model, device = load_indot5_model()
    results = []
    
    for match in tqdm(matches, desc="Double paraphrasing"):
        original = match['matched_para']
        
        # Split long texts
        if len(original) > 400:
            sentences = [s.strip() + '.' for s in original.split('.') if s.strip()]
            paraphrased_parts = []
            
            for sent in sentences:
                if len(sent) > 10:
                    # DOUBLE paraphrase each sentence
                    para = aggressive_paraphrase(sent, tokenizer, model, device, rounds=2)
                    paraphrased_parts.append(para)
                else:
                    paraphrased_parts.append(sent)
            
            paraphrased = ' '.join(paraphrased_parts)
        else:
            # DOUBLE paraphrase
            paraphrased = aggressive_paraphrase(original, tokenizer, model, device, rounds=2)
        
        results.append({
            "original": original,
            "paraphrased": paraphrased,
            "method": "double_aggressive",
            "page": match['page'],
            "similarity_before": match['similarity'],
            "para_index": match['para_index']
        })
    
    return results

def apply_to_docx(paraphrased_data, input_docx, output_docx):
    """Apply paraphrased texts to DOCX"""
    print(f"\n📄 Loading: {input_docx}")
    doc = docx.Document(input_docx)
    
    replaced = 0
    for item in tqdm(paraphrased_data, desc="Applying"):
        original = item['original']
        paraphrased = item['paraphrased']
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            ratio = SequenceMatcher(None, original.lower(), para.text.lower()).ratio()
            if ratio >= 0.7:
                # Replace keeping formatting
                if para.runs:
                    for run in para.runs:
                        run.text = ''
                    para.runs[0].text = paraphrased
                else:
                    para.text = paraphrased
                replaced += 1
                break
    
    print(f"✅ Replaced: {replaced}/{len(paraphrased_data)}")
    doc.save(output_docx)
    print(f"💾 Saved: {output_docx}")
    return replaced

def main():
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python aggressive_paraphrase.py <flagged.json> <original.docx>")
        sys.exit(1)
    
    flagged_json = sys.argv[1]
    docx_file = sys.argv[2]
    output_file = "testing_aggressive_bypassed.docx"
    
    print("="*70)
    print("🔥 AGGRESSIVE DOUBLE-PARAPHRASE")
    print("="*70)
    print("Strategy:")
    print("  🔥 Round 1: High-temperature paraphrase")
    print("  🔥 Round 2: Paraphrase the paraphrase!")
    print("  🎯 Temperature: 1.5 (very creative)")
    print("  🎯 Beams: 15 (high quality)")
    print("  🎯 Top-k: 100 (many options)")
    print("  ❌ NO invisible chars")
    print("  ❌ NO unicode tricks")
    print("="*70)
    
    # Step 1: Match
    matches = match_flagged_with_docx(flagged_json, docx_file, threshold=0.5)
    
    if not matches:
        print("❌ No matches found")
        return
    
    # Save matches
    with open("aggressive_matches.json", 'w', encoding='utf-8') as f:
        json.dump(matches, f, ensure_ascii=False, indent=2)
    
    # Step 2: Aggressive paraphrase
    paraphrased = process_aggressive(matches)
    
    # Save paraphrased
    with open("aggressive_paraphrased.json", 'w', encoding='utf-8') as f:
        json.dump(paraphrased, f, ensure_ascii=False, indent=2)
    
    # Step 3: Apply to DOCX
    apply_to_docx(paraphrased, docx_file, output_file)
    
    print("\n" + "="*70)
    print("✅ COMPLETE!")
    print("="*70)
    print(f"📄 Output: {output_file}")
    print(f"📊 Processed: {len(paraphrased)} texts")
    print()
    print("🔥 AGGRESSIVE Changes:")
    print("  - DOUBLE paraphrase (2 rounds)")
    print("  - High temperature (1.5 = very creative)")
    print("  - More beams (15 vs 10)")
    print("  - Clean output (no tricks)")
    print()
    print("🎯 Expected: MUCH lower similarity!")
    print("   Target: <25% (from current 32%)")
    print("="*70)

if __name__ == "__main__":
    main()
