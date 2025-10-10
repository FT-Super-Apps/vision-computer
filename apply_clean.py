#!/usr/bin/env python3
"""
Apply CLEAN paraphrased texts to DOCX
No invisible chars, no unicode substitution - just pure paraphrase
"""

import json
import docx
from difflib import SequenceMatcher
from pathlib import Path

def find_and_replace_paragraph(doc, original_text, new_text, threshold=0.7):
    """
    Find paragraph with fuzzy matching and replace it
    
    Args:
        doc: Document object
        original_text: Original text to find
        new_text: New text to replace with
        threshold: Minimum similarity to consider a match
    
    Returns:
        (bool, str): (success, message)
    """
    best_match = None
    best_ratio = 0
    best_para_idx = -1
    
    for idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        if not para_text:
            continue
        
        ratio = SequenceMatcher(None, original_text.lower(), para_text.lower()).ratio()
        
        if ratio > best_ratio:
            best_ratio = ratio
            best_match = para
            best_para_idx = idx
    
    if best_ratio >= threshold and best_match is not None:
        # Replace paragraph text while preserving formatting
        for run in best_match.runs:
            run.text = ''
        
        if best_match.runs:
            best_match.runs[0].text = new_text
        else:
            best_match.add_run(new_text)
        
        return True, f"Replaced (similarity: {best_ratio*100:.1f}%)"
    
    return False, f"Not found (best match: {best_ratio*100:.1f}%)"

def apply_paraphrased_to_docx(paraphrased_json, docx_path, output_path):
    """Apply paraphrased texts to DOCX"""
    
    with open(paraphrased_json, 'r', encoding='utf-8') as f:
        results = json.load(f)
    
    doc = docx.Document(docx_path)
    
    print(f"\n📝 Applying {len(results)} paraphrased texts to DOCX...")
    print("="*70)
    
    success_count = 0
    failed = []
    
    for i, result in enumerate(results, 1):
        original = result['original']
        paraphrased = result['paraphrased']
        method = result.get('method', 'unknown')
        
        success, msg = find_and_replace_paragraph(doc, original, paraphrased, threshold=0.7)
        
        status = "✅" if success else "❌"
        print(f"{i}. {status} Page {result['page']} ({method}): {msg}")
        
        if success:
            success_count += 1
        else:
            failed.append({
                'index': i,
                'original': original[:80] + '...',
                'reason': msg
            })
    
    # Save modified document
    doc.save(output_path)
    
    print("\n" + "="*70)
    print("📊 SUMMARY")
    print("="*70)
    print(f"   Total texts    : {len(results)}")
    print(f"   Applied        : {success_count}")
    print(f"   Failed         : {len(failed)}")
    print(f"   Success rate   : {success_count/len(results)*100:.1f}%")
    print(f"\n💾 Output saved  : {output_path}")
    
    if failed:
        print(f"\n⚠️  Failed replacements:")
        for item in failed:
            print(f"   {item['index']}. {item['original']} - {item['reason']}")
    
    print("="*70)

def main():
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python apply_clean.py <paraphrased.json> <original.docx>")
        sys.exit(1)
    
    paraphrased_json = sys.argv[1]
    docx_path = sys.argv[2]
    
    # Generate output filename
    stem = Path(docx_path).stem
    output_path = f"{stem}_clean_bypassed.docx"
    
    print("="*70)
    print("🎓 APPLY CLEAN PARAPHRASED TEXTS")
    print("="*70)
    
    apply_paraphrased_to_docx(paraphrased_json, docx_path, output_path)

if __name__ == "__main__":
    main()
