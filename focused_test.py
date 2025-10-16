#!/usr/bin/env python3
"""
FOCUSED TEST: Process only specific pages
Quick iteration to find best paraphrase strategy
"""

import json
import docx
from difflib import SequenceMatcher
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch
from tqdm import tqdm
import sys

def load_indot5_model(model_name="Wikidepia/IndoT5-base-paraphrase"):
    """Load IndoT5 model"""
    print(f"🔧 Loading IndoT5: {model_name}")
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)
    print(f"✅ Model loaded on {device}")
    return tokenizer, model, device

def ultra_aggressive_paraphrase(text, tokenizer, model, device, rounds=3):
    """
    ULTRA AGGRESSIVE: 3 rounds + extreme temperature
    """
    current_text = text
    
    for round_num in range(1, rounds + 1):
        input_text = f"paraphrase: {current_text}"
        inputs = tokenizer(
            input_text,
            return_tensors="pt",
            max_length=512,
            truncation=True,
            padding=True
        ).to(device)
        
        # ULTRA aggressive parameters
        temperature = 1.5 + (round_num * 0.2)  # Increase each round
        
        with torch.no_grad():
            outputs = model.generate(
                inputs.input_ids,
                max_length=512,
                num_beams=20,              # Max beams
                temperature=temperature,    # Progressive temperature
                do_sample=True,
                top_k=150,                 # Very high
                top_p=0.99,                # Almost all
                early_stopping=True,
                no_repeat_ngram_size=5,
                num_return_sequences=1
            )
        
        current_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
    
    return current_text

def filter_by_pages(flagged_json, target_pages):
    """Filter flagged texts to only specific pages"""
    with open(flagged_json, 'r', encoding='utf-8') as f:
        flagged = json.load(f)
    
    filtered = [item for item in flagged if item['page'] in target_pages]
    return filtered

def match_flagged_with_docx(flagged_texts, docx_path, threshold=0.5):
    """Match flagged texts with DOCX"""
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    print(f"\n🔍 Matching {len(flagged_texts)} flagged with {len(paragraphs)} paragraphs...")
    
    matches = []
    for i, flagged in enumerate(flagged_texts, 1):
        flagged_text = flagged['text']
        best_match = None
        best_ratio = 0
        best_idx = -1
        
        for idx, para in enumerate(paragraphs):
            ratio = SequenceMatcher(None, flagged_text.lower(), para.lower()).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = para
                best_idx = idx
        
        if best_ratio < 0.5:
            keywords = [w.lower() for w in flagged_text.split() if len(w) > 4]
            for idx, para in enumerate(paragraphs):
                keyword_matches = sum(1 for kw in keywords if kw in para.lower())
                keyword_ratio = keyword_matches / len(keywords) if keywords else 0
                if keyword_ratio > best_ratio:
                    best_ratio = keyword_ratio
                    best_match = para
                    best_idx = idx
        
        if best_ratio >= threshold:
            matches.append({
                "flagged_text": flagged_text,
                "matched_para": best_match,
                "similarity": round(best_ratio * 100, 2),
                "para_index": best_idx,
                "page": flagged['page']
            })
    
    print(f"✅ Matched: {len(matches)}/{len(flagged_texts)}")
    return matches

def process_ultra_aggressive(matches, rounds=3):
    """ULTRA AGGRESSIVE paraphrase"""
    print(f"\n🔥🔥🔥 ULTRA AGGRESSIVE: {len(matches)} texts, {rounds} rounds...")
    print("="*70)
    
    tokenizer, model, device = load_indot5_model()
    results = []
    
    for match in tqdm(matches, desc=f"Paraphrasing {rounds}x"):
        original = match['matched_para']
        
        # Split long texts
        if len(original) > 400:
            sentences = [s.strip() + '.' for s in original.split('.') if s.strip()]
            paraphrased_parts = []
            
            for sent in sentences:
                if len(sent) > 10:
                    para = ultra_aggressive_paraphrase(sent, tokenizer, model, device, rounds=rounds)
                    paraphrased_parts.append(para)
                else:
                    paraphrased_parts.append(sent)
            
            paraphrased = ' '.join(paraphrased_parts)
        else:
            paraphrased = ultra_aggressive_paraphrase(original, tokenizer, model, device, rounds=rounds)
        
        results.append({
            "original": original,
            "paraphrased": paraphrased,
            "method": f"ultra_aggressive_{rounds}x",
            "page": match['page'],
            "similarity_before": match['similarity'],
            "para_index": match['para_index']
        })
    
    return results

def apply_to_docx(paraphrased_data, input_docx, output_docx):
    """Apply paraphrased texts to DOCX"""
    print(f"\n📄 Loading: {input_docx}")
    doc = docx.Document(input_docx)
    
    replaced = 0
    for item in tqdm(paraphrased_data, desc="Applying"):
        original = item['original']
        paraphrased = item['paraphrased']
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            ratio = SequenceMatcher(None, original.lower(), para.text.lower()).ratio()
            if ratio >= 0.7:
                if para.runs:
                    for run in para.runs:
                        run.text = ''
                    para.runs[0].text = paraphrased
                else:
                    para.text = paraphrased
                replaced += 1
                break
    
    print(f"✅ Replaced: {replaced}/{len(paraphrased_data)}")
    doc.save(output_docx)
    print(f"💾 Saved: {output_docx}")
    return replaced

def main():
    if len(sys.argv) < 4:
        print("Usage: python focused_test.py <flagged.json> <original.docx> <pages>")
        print("Example: python focused_test.py testing_flagged.json testing.docx 24,14")
        sys.exit(1)
    
    flagged_json = sys.argv[1]
    docx_file = sys.argv[2]
    pages_str = sys.argv[3]
    
    # Parse pages
    target_pages = [int(p.strip()) for p in pages_str.split(',')]
    
    output_file = f"testing_focused_pages_{pages_str.replace(',', '_')}.docx"
    
    print("="*70)
    print("🎯 FOCUSED TEST: Specific Pages Only")
    print("="*70)
    print(f"Target pages: {target_pages}")
    print()
    print("Strategy:")
    print("  🔥 3-ROUND paraphrase")
    print("  🔥 Temperature: 1.5 → 1.7 → 1.9 (progressive)")
    print("  🔥 Beams: 20 (maximum quality)")
    print("  🔥 Top-k: 150, Top-p: 0.99 (extreme)")
    print("  ❌ NO invisible chars/unicode")
    print("="*70)
    
    # Step 1: Filter by pages
    print(f"\n📄 Filtering flagged texts for pages {target_pages}...")
    filtered_texts = filter_by_pages(flagged_json, target_pages)
    print(f"✅ Found {len(filtered_texts)} texts in target pages")
    
    if not filtered_texts:
        print("❌ No texts found in target pages")
        return
    
    # Show what we're processing
    print(f"\n📋 Texts to process:")
    for i, item in enumerate(filtered_texts, 1):
        print(f"{i}. [Page {item['page']}] {item['text'][:60]}...")
    
    # Step 2: Match with DOCX
    matches = match_flagged_with_docx(filtered_texts, docx_file, threshold=0.5)
    
    if not matches:
        print("❌ No matches found")
        return
    
    # Save matches
    with open("focused_matches.json", 'w', encoding='utf-8') as f:
        json.dump(matches, f, ensure_ascii=False, indent=2)
    
    # Step 3: ULTRA aggressive paraphrase
    paraphrased = process_ultra_aggressive(matches, rounds=3)
    
    # Save paraphrased
    with open("focused_paraphrased.json", 'w', encoding='utf-8') as f:
        json.dump(paraphrased, f, ensure_ascii=False, indent=2)
    
    # Step 4: Apply to DOCX
    apply_to_docx(paraphrased, docx_file, output_file)
    
    print("\n" + "="*70)
    print("✅ FOCUSED TEST COMPLETE!")
    print("="*70)
    print(f"📄 Output: {output_file}")
    print(f"📊 Pages processed: {target_pages}")
    print(f"📊 Texts modified: {len(paraphrased)}")
    print()
    print("🎯 Next Steps:")
    print("  1. Upload to Turnitin")
    print("  2. Check ONLY those pages")
    print("  3. Compare similarity change")
    print("  4. If works → apply to all pages!")
    print("="*70)
    
    # Show sample
    if paraphrased:
        print("\n📝 SAMPLE COMPARISON:")
        print("="*70)
        sample = paraphrased[0]
        print(f"ORIGINAL:")
        print(f"  {sample['original'][:150]}...")
        print()
        print(f"ULTRA PARAPHRASE (3 rounds):")
        print(f"  {sample['paraphrased'][:150]}...")
        print("="*70)

if __name__ == "__main__":
    main()
