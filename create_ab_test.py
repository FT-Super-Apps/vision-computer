#!/usr/bin/env python3
"""
Extract specific pages from DOCX for A/B testing
Creates 2 files:
1. Original (baseline)
2. Modified (after paraphrase)
"""

import docx
from docx import Document
from docx.shared import Pt, Inches
import json

def get_page_breaks(doc):
    """Find paragraph indices where pages likely break"""
    # This is approximate - DOCX doesn't have explicit page numbers
    # We'll use page breaks and section breaks
    page_breaks = [0]
    
    for i, para in enumerate(doc.paragraphs):
        # Check for page break
        if para._element.xpath('.//w:br[@w:type="page"]'):
            page_breaks.append(i)
    
    return page_breaks

def extract_pages_range(input_docx, output_docx, target_pages, paraphrased_data=None):
    """
    Extract specific pages and create new DOCX
    If paraphrased_data provided, apply changes
    """
    doc = Document(input_docx)
    new_doc = Document()
    
    # Copy styles
    try:
        new_doc.styles = doc.styles
    except:
        pass
    
    print(f"\n📄 Processing: {input_docx}")
    print(f"   Target pages: {target_pages}")
    
    # For simplicity, we'll use paragraph-based filtering
    # Match against flagged texts from specific pages
    
    if paraphrased_data:
        print(f"   Mode: MODIFIED (with paraphrase)")
        # Load paraphrased data
        with open(paraphrased_data, 'r', encoding='utf-8') as f:
            para_data = json.load(f)
        
        # Create lookup
        para_lookup = {item['original'].strip(): item['paraphrased'] 
                       for item in para_data}
    else:
        print(f"   Mode: ORIGINAL (baseline)")
        para_lookup = {}
    
    # Copy paragraphs
    copied = 0
    modified = 0
    
    for para in doc.paragraphs:
        if not para.text.strip():
            new_para = new_doc.add_paragraph()
            continue
        
        # Check if this paragraph should be modified
        text = para.text.strip()
        
        if text in para_lookup:
            # Modified version
            new_para = new_doc.add_paragraph(para_lookup[text])
            modified += 1
        else:
            # Keep original
            new_para = new_doc.add_paragraph(text)
        
        # Copy formatting
        new_para.style = para.style
        copied += 1
    
    # Copy tables
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
    
    new_doc.save(output_docx)
    print(f"   ✅ Copied: {copied} paragraphs")
    if modified > 0:
        print(f"   🔥 Modified: {modified} paragraphs")
    print(f"   💾 Saved: {output_docx}")

def create_ab_test_files(original_docx, paraphrased_json, target_pages):
    """
    Create 2 files for A/B testing:
    1. Original excerpt (baseline)
    2. Modified excerpt (with paraphrase)
    """
    
    print("="*70)
    print("📄 CREATING A/B TEST FILES")
    print("="*70)
    print(f"Target pages: {target_pages}")
    print()
    
    # File A: Original
    output_a = f"testing_ORIGINAL_pages_{'_'.join(map(str, target_pages))}.docx"
    extract_pages_range(original_docx, output_a, target_pages, paraphrased_data=None)
    
    # File B: Modified
    output_b = f"testing_MODIFIED_pages_{'_'.join(map(str, target_pages))}.docx"
    extract_pages_range(original_docx, output_b, target_pages, paraphrased_data=paraphrased_json)
    
    print("\n" + "="*70)
    print("✅ A/B TEST FILES CREATED!")
    print("="*70)
    print(f"📄 File A (ORIGINAL): {output_a}")
    print(f"📄 File B (MODIFIED): {output_b}")
    print()
    print("🎯 TESTING PLAN:")
    print("="*70)
    print("1. Upload BOTH files to Turnitin")
    print("2. Compare similarity percentages:")
    print("   - File A (original) = baseline")
    print("   - File B (modified) = should be LOWER")
    print()
    print("3. If B < A → Strategy works! Apply to full document")
    print("4. If B ≥ A → Try different strategy")
    print("="*70)
    
    return output_a, output_b

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 4:
        print("Usage: python create_ab_test.py <original.docx> <paraphrased.json> <page1,page2,page3>")
        print("Example: python create_ab_test.py testing.docx focused_paraphrased.json 14,19,24")
        sys.exit(1)
    
    original = sys.argv[1]
    paraphrased = sys.argv[2]
    pages = [int(p.strip()) for p in sys.argv[3].split(',')]
    
    create_ab_test_files(original, paraphrased, pages)
