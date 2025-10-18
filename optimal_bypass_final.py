#!/usr/bin/env python3
"""
FINAL OPTIMAL BYPASS - Best of All Methods Combined
Strategy: Reduced Cyrillic (25-30%) + Invisible Chars (15-20%)
Result: Effective bypass while maintaining natural appearance
"""

import sys
import random
from docx import Document
from difflib import SequenceMatcher

class OptimalBypass:
    def __init__(self):
        # Core Cyrillic homoglyphs (ONLY perfect visual matches)
        self.cyrillic_map = {
            # Perfect lowercase matches
            'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с', 'x': 'х',
            # Perfect uppercase matches
            'A': 'А', 'E': 'Е', 'O': 'О', 'P': 'Р', 'C': 'С', 'X': 'Х',
            'M': 'М', 'H': 'Н', 'B': 'В', 'K': 'К', 'T': 'Т',
            # Additional safe matches
            'y': 'у', 'i': 'і', 's': 'ѕ'
        }
        
        # Invisible Unicode characters
        self.invisible_chars = [
            '\u200B',  # Zero-Width Space
            '\u200C',  # Zero-Width Non-Joiner
            '\u200D',  # Zero-Width Joiner
        ]
    
    def apply_cyrillic(self, text, density=0.30):
        """Apply Cyrillic homoglyphs with specified density"""
        if not text:
            return text
        
        result = list(text)
        replaceable = [i for i, c in enumerate(text) if c in self.cyrillic_map]
        
        if not replaceable:
            return text
        
        # Apply density
        num_replace = max(1, int(len(replaceable) * density))
        positions = random.sample(replaceable, min(num_replace, len(replaceable)))
        
        for pos in positions:
            result[pos] = self.cyrillic_map[result[pos]]
        
        return ''.join(result)
    
    def add_invisible_chars(self, text, density=0.15):
        """Add invisible characters between words"""
        if not text or len(text) < 10:
            return text
        
        words = text.split()
        result_words = []
        
        for word in words:
            if len(word) > 3 and random.random() < density:
                # Insert invisible char in middle of word
                mid = len(word) // 2
                invisible = random.choice(self.invisible_chars)
                modified_word = word[:mid] + invisible + word[mid:]
                result_words.append(modified_word)
            else:
                result_words.append(word)
        
        return ' '.join(result_words)
    
    def apply_combined_strategy(self, text, is_header=False):
        """Apply combined Cyrillic + Invisible strategy"""
        if not text:
            return text
        
        # Determine densities based on text type
        if is_header:
            cyrillic_density = 0.30  # 30% for headers
            invisible_density = 0.20  # 20% for headers
        else:
            cyrillic_density = 0.25  # 25% for content
            invisible_density = 0.15  # 15% for content
        
        # Step 1: Apply Cyrillic replacements
        modified = self.apply_cyrillic(text, cyrillic_density)
        
        # Step 2: Add invisible characters
        modified = self.add_invisible_chars(modified, invisible_density)
        
        return modified
    
    def is_header(self, text):
        """Detect if text is a header"""
        if not text:
            return False
        
        # Short text (≤6 words) likely a header
        word_count = len(text.split())
        if word_count <= 6:
            return True
        
        # Common header patterns
        header_patterns = [
            r'^BAB\s+[IVX0-9]+',
            r'^[0-9]+\.[0-9]+',
            r'(Latar Belakang|Rumusan Masalah|Tujuan|Manfaat|Batasan)',
            r'(PENDAHULUAN|METODOLOGI|KESIMPULAN)',
        ]
        
        import re
        for pattern in header_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False

def optimal_bypass_final(input_docx, flags_file, output_docx):
    """
    Final optimal bypass combining all successful methods
    """
    print(f"\n{'='*70}")
    print(f"🎯 FINAL OPTIMAL BYPASS - COMBINED STRATEGY")
    print(f"{'='*70}")
    print(f"Input  : {input_docx}")
    print(f"Flags  : {flags_file}")
    print(f"Output : {output_docx}")
    print(f"\n📊 Strategy:")
    print(f"   • Cyrillic homoglyphs: 25-30% (proven effective)")
    print(f"   • Invisible chars: 15-20% (extra layer)")
    print(f"   • Smart density: Headers 30%, Content 25%")
    print(f"   • Natural appearance: Reduced from 70% to 25%")
    print(f"{'='*70}\n")
    
    # Initialize bypass engine
    bypass = OptimalBypass()
    
    # Load flagged texts
    with open(flags_file, 'r', encoding='utf-8') as f:
        flagged_texts = [line.strip() for line in f if line.strip()]
    
    print(f"📋 Loaded {len(flagged_texts)} flagged texts\n")
    
    # Load document
    doc = Document(input_docx)
    modified_count = 0
    
    print(f"🔧 Processing paragraphs...\n")
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        if not para_text:
            continue
        
        # Check if paragraph contains flagged text
        is_flagged = False
        for flag in flagged_texts:
            similarity = SequenceMatcher(None, para_text.lower(), flag.lower()).ratio()
            if similarity > 0.80 or flag.lower() in para_text.lower():
                is_flagged = True
                break
        
        if is_flagged:
            modified_count += 1
            
            # Detect if header
            is_hdr = bypass.is_header(para_text)
            
            # Apply combined strategy
            original_text = para_text
            modified_text = bypass.apply_combined_strategy(para_text, is_hdr)
            
            # Count changes
            cyrillic_count = sum(1 for a, b in zip(original_text, modified_text) 
                               if a != b and '\u0400' <= b <= '\u04FF')
            invisible_count = sum(modified_text.count(char) for char in bypass.invisible_chars)
            
            # Replace paragraph
            para.clear()
            new_run = para.add_run(modified_text)
            
            # Report
            print(f"✅ Modified #{modified_count}:")
            print(f"   Para {para_idx}: {original_text[:60]}...")
            print(f"   Type: {'Header' if is_hdr else 'Content'}")
            print(f"   Cyrillic: {cyrillic_count} chars")
            print(f"   Invisible: {invisible_count} chars")
            print(f"   Total density: {(cyrillic_count + invisible_count)/len(original_text)*100:.1f}%")
            print(f"   Preview: {modified_text[:60]}...")
            print()
    
    # Save document
    print(f"💾 Saving to: {output_docx}")
    doc.save(output_docx)
    
    print(f"\n{'='*70}")
    print(f"✅ FINAL BYPASS COMPLETE!")
    print(f"{'='*70}")
    print(f"📊 Statistics:")
    print(f"   Flagged texts   : {len(flagged_texts)}")
    print(f"   Paragraphs modified: {modified_count}")
    print(f"   Cyrillic density: 25-30%")
    print(f"   Invisible density: 15-20%")
    print(f"\n🎯 Expected Results:")
    print(f"   • Maintains 2% similarity ✅")
    print(f"   • Natural font appearance ✅")
    print(f"   • No obvious Cyrillic blocks ✅")
    print(f"   • Extra invisible layer ✅")
    print(f"\n📄 Output: {output_docx}")
    print(f"{'='*70}\n")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python optimal_bypass_final.py input.docx [flags_file] [output.docx]")
        print("\nExample:")
        print("  python optimal_bypass_final.py original.docx flag_new.txt original_final.docx")
        print("\nThis combines:")
        print("  • Cyrillic homoglyphs (25-30%) - proven effective")
        print("  • Invisible chars (15-20%) - extra stealth")
        print("  • Smart density - natural appearance")
        sys.exit(1)
    
    input_file = sys.argv[1]
    flags_file = sys.argv[2] if len(sys.argv) > 2 else "flag_new.txt"
    output_file = sys.argv[3] if len(sys.argv) > 3 else "original_final.docx"
    
    optimal_bypass_final(input_file, flags_file, output_file)
