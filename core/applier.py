"""
Core module untuk apply paraphrased text ke DOCX
"""

from docx import Document
from difflib import SequenceMatcher
from typing import List, Dict

def find_and_replace_paragraph(doc: Document, original_text: str, new_text: str, 
                               similarity_threshold: float = 0.7) -> bool:
    """Find and replace paragraph dengan similarity check"""
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue
        
        similarity = SequenceMatcher(None, original_text.lower(), para_text.lower()).ratio()
        
        if similarity >= similarity_threshold:
            # Replace entire paragraph
            para.clear()
            para.add_run(new_text)
            return True
    
    return False

def apply_paraphrased_to_docx(paraphrased_results: List[Dict], input_docx: str, 
                              output_docx: str) -> Dict:
    """Apply all paraphrased texts to DOCX"""
    doc = Document(input_docx)
    
    stats = {
        'total': len(paraphrased_results),
        'applied': 0,
        'skipped': 0,
        'failed': []
    }
    
    for item in paraphrased_results:
        original = item['matched_text']
        bypassed = item['bypassed_text']
        
        if find_and_replace_paragraph(doc, original, bypassed):
            stats['applied'] += 1
        else:
            stats['skipped'] += 1
            stats['failed'].append({
                'original': original[:50] + '...',
                'reason': 'Not found in document'
            })
    
    doc.save(output_docx)
    return stats
