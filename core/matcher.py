"""
Core module untuk matching flagged text dengan DOCX
"""

from difflib import SequenceMatcher
from docx import Document
from typing import List, Dict

def match_flagged_with_docx(flagged_texts: List[Dict], docx_path: str, threshold: float = 0.5) -> List[Dict]:
    """Match flagged texts dengan paragraf dalam DOCX"""
    doc = Document(docx_path)
    matches = []
    
    for flagged in flagged_texts:
        flagged_text = flagged['text']
        best_match = None
        best_similarity = 0
        
        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if not para_text:
                continue
            
            # Similarity matching
            similarity = SequenceMatcher(None, flagged_text.lower(), para_text.lower()).ratio()
            
            if similarity > best_similarity:
                best_similarity = similarity
                best_match = {
                    'paragraph_index': para_idx,
                    'original_text': para_text,
                    'similarity': similarity
                }
        
        # Jika similarity terlalu rendah, coba keyword matching
        if best_similarity < threshold:
            flagged_words = set(flagged_text.lower().split())
            for para_idx, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                if not para_text:
                    continue
                
                para_words = set(para_text.lower().split())
                common_words = flagged_words & para_words
                keyword_score = len(common_words) / len(flagged_words) if flagged_words else 0
                
                if keyword_score > best_similarity:
                    best_similarity = keyword_score
                    best_match = {
                        'paragraph_index': para_idx,
                        'original_text': para_text,
                        'similarity': keyword_score
                    }
        
        if best_match and best_similarity >= threshold:
            matches.append({
                'flagged_text': flagged_text,
                'page': flagged.get('page'),
                'matched_text': best_match['original_text'],
                'paragraph_index': best_match['paragraph_index'],
                'similarity': best_similarity
            })
    
    return matches
