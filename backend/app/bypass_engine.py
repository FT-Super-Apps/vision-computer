#!/usr/bin/env python3
"""
Bypass Engine Core
Integrates all bypass strategies
"""

import random
import re
from docx import Document
from datetime import datetime
from pathlib import Path
import sys

# Add parent directory to path
sys.path.append(str(Path(__file__).parent.parent))

from config import (
    HOMOGLYPHS, INVISIBLE_CHARS,
    HEADER_PATTERNS, KNOWN_HEADERS, STANDARD_PHRASES,
    TARGETED_CONFIG, TARGETED_AGGRESSIVE_CONFIG, HEADER_CONFIG,
    DEFAULT_CONFIG
)

class BypassEngine:
    """Main bypass engine"""

    def __init__(self):
        self.homoglyphs = HOMOGLYPHS
        self.invisible_chars = INVISIBLE_CHARS
        self.header_patterns = HEADER_PATTERNS
        self.known_headers = KNOWN_HEADERS
        self.standard_phrases = STANDARD_PHRASES

        self.strategies = {
            'natural': TARGETED_CONFIG,
            'aggressive': TARGETED_AGGRESSIVE_CONFIG,
            'header_focused': HEADER_CONFIG
        }

    # ========================================================================
    # BYPASS TECHNIQUES
    # ========================================================================

    def apply_homoglyphs(self, text: str, density: float = 0.50) -> str:
        """Apply homoglyphs with smart selection"""
        if not text:
            return text

        result = list(text)
        replaceable = [i for i, c in enumerate(text) if c in self.homoglyphs]

        if not replaceable:
            return text

        # Smart selection: prioritize natural-looking characters
        high_priority = {'a', 'e', 'o', 'c', 'p', 'x', 'A', 'E', 'O', 'C', 'P', 'X'}
        high_priority_pos = [i for i in replaceable if result[i] in high_priority]
        other_pos = [i for i in replaceable if result[i] not in high_priority]

        # Replace more from high priority
        num_high = int(len(high_priority_pos) * min(density * 1.2, 1.0))
        num_other = int(len(other_pos) * density * 0.5)

        selected_pos = []
        if high_priority_pos:
            selected_pos.extend(random.sample(high_priority_pos, min(num_high, len(high_priority_pos))))
        if other_pos and num_other > 0:
            selected_pos.extend(random.sample(other_pos, min(num_other, len(other_pos))))

        for pos in selected_pos:
            result[pos] = self.homoglyphs[result[pos]]

        return ''.join(result)

    def apply_invisible_chars(self, text: str, density: float = 0.15) -> str:
        """Insert invisible characters at word boundaries"""
        if not text or len(text) < 3:
            return text

        result = list(text)
        word_boundaries = [i for i, c in enumerate(text) if c == ' ']

        if not word_boundaries:
            return text

        num_insert = max(1, int(len(word_boundaries) * density))
        selected_boundaries = random.sample(word_boundaries, min(num_insert, len(word_boundaries)))

        for pos in sorted(selected_boundaries, reverse=True):
            if pos + 1 < len(result):
                invisible_char = random.choice(self.invisible_chars)
                result.insert(pos + 1, invisible_char)

        return ''.join(result)

    def apply_combined_bypass(self, text: str, homoglyph_density: float, invisible_density: float) -> str:
        """Combined bypass technique"""
        text = self.apply_homoglyphs(text, density=homoglyph_density)
        text = self.apply_invisible_chars(text, density=invisible_density)
        return text

    # ========================================================================
    # DETECTION HELPERS
    # ========================================================================

    def is_header(self, text: str) -> bool:
        """Check if text is a header"""
        text = text.strip()

        for header in self.known_headers:
            if header.lower() in text.lower():
                return True

        for pattern in self.header_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True

        return False

    def is_standard_phrase(self, text: str) -> bool:
        """Check if text contains standard academic phrases"""
        text_lower = text.lower()
        for phrase in self.standard_phrases:
            if phrase.lower() in text_lower:
                return True
        return False

    # ========================================================================
    # MAIN PROCESSING
    # ========================================================================

    def process_bypass(self, input_path: str, strategy: str = 'header_focused',
                      homoglyph_density: float = None, invisible_density: float = None,
                      output_path: str = None) -> dict:
        """
        Main bypass processing function

        Args:
            input_path: Path to input document
            strategy: Bypass strategy (natural, aggressive, header_focused)
            homoglyph_density: Custom homoglyph density (overrides strategy)
            invisible_density: Custom invisible density (overrides strategy)
            output_path: Custom output path

        Returns:
            dict with processing results
        """

        # Get strategy config
        if strategy not in self.strategies:
            strategy = 'header_focused'

        config = self.strategies[strategy]

        # Use custom densities if provided
        h_density = homoglyph_density if homoglyph_density is not None else config['homoglyph_density']
        i_density = invisible_density if invisible_density is not None else config['invisible_density']

        # Generate output path
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            input_filename = Path(input_path).stem
            output_path = f"outputs/{input_filename}_bypassed_{timestamp}.docx"

        # Load document
        doc = Document(input_path)

        header_count = 0
        phrase_count = 0
        total_modifications = 0

        # Process each paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            modified = False
            is_hdr = self.is_header(text)
            is_std = self.is_standard_phrase(text)

            # Apply bypass based on strategy
            if strategy == 'header_focused':
                # Only process headers and standard phrases
                if is_hdr or is_std:
                    for run in para.runs:
                        if run.text:
                            original_text = run.text
                            modified_text = self.apply_combined_bypass(
                                original_text, h_density, i_density
                            )
                            if modified_text != original_text:
                                run.text = modified_text
                                modified = True

                    if modified:
                        total_modifications += 1
                        if is_hdr:
                            header_count += 1
                        if is_std:
                            phrase_count += 1

            else:
                # Process all content (natural/aggressive)
                for run in para.runs:
                    if run.text:
                        original_text = run.text
                        modified_text = self.apply_combined_bypass(
                            original_text, h_density, i_density
                        )
                        if modified_text != original_text:
                            run.text = modified_text
                            modified = True

                if modified:
                    total_modifications += 1

        # Save document
        doc.save(output_path)

        return {
            'success': True,
            'message': 'Document processed successfully',
            'input_file': input_path,
            'output_file': output_path,
            'strategy': strategy,
            'statistics': {
                'headers_modified': header_count,
                'phrases_modified': phrase_count,
                'total_modifications': total_modifications,
                'homoglyph_density': h_density,
                'invisible_density': i_density
            }
        }

    # ========================================================================
    # ANALYSIS FUNCTIONS
    # ========================================================================

    def analyze_document(self, input_path: str, flag_file: str = 'flag.txt') -> dict:
        """Analyze document for flagged phrases"""

        doc = Document(input_path)
        full_text = '\n'.join([para.text for para in doc.paragraphs])

        flagged_phrases = []

        # Load flag file if exists
        if Path(flag_file).exists():
            with open(flag_file, 'r', encoding='utf-8') as f:
                flags = [line.strip() for line in f if line.strip()]

            for phrase in flags:
                count = full_text.count(phrase)
                if count > 0:
                    flagged_phrases.append({
                        'phrase': phrase[:100],
                        'category': 'flagged',
                        'occurrences': count
                    })

        # Check for headers
        header_count = 0
        phrase_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if self.is_header(text):
                header_count += 1
                flagged_phrases.append({
                    'phrase': text[:100],
                    'category': 'header',
                    'occurrences': 1
                })
            elif self.is_standard_phrase(text):
                phrase_count += 1

        return {
            'success': True,
            'message': 'Analysis completed',
            'filename': input_path,
            'total_paragraphs': len(doc.paragraphs),
            'flagged_phrases': flagged_phrases,
            'statistics': {
                'total_flags': len(flagged_phrases),
                'headers': header_count,
                'standard_phrases': phrase_count
            }
        }

    def compare_documents(self, original_path: str, bypassed_path: str, flag_file: str = 'flag.txt') -> dict:
        """Compare original vs bypassed document"""

        doc_original = Document(original_path)
        doc_bypassed = Document(bypassed_path)

        text_original = '\n'.join([para.text for para in doc_original.paragraphs])
        text_bypassed = '\n'.join([para.text for para in doc_bypassed.paragraphs])

        # Load flags
        flags = []
        if Path(flag_file).exists():
            with open(flag_file, 'r', encoding='utf-8') as f:
                flags = [line.strip() for line in f if line.strip()]

        hidden_count = 0
        still_detected = 0
        comparison_details = []

        for phrase in flags:
            count_orig = text_original.count(phrase)
            count_bypass = text_bypassed.count(phrase)

            if count_orig > 0:
                status = "hidden" if count_bypass == 0 else "detected"

                if count_bypass == 0:
                    hidden_count += 1
                else:
                    still_detected += 1

                comparison_details.append({
                    'phrase': phrase[:50],
                    'original_count': count_orig,
                    'bypassed_count': count_bypass,
                    'status': status
                })

        total = hidden_count + still_detected
        success_rate = (hidden_count / total * 100) if total > 0 else 0

        return {
            'success': True,
            'message': 'Comparison completed',
            'original_file': original_path,
            'bypassed_file': bypassed_path,
            'comparison': {
                'total_flags': total,
                'hidden_flags': hidden_count,
                'still_detected': still_detected,
                'details': comparison_details
            },
            'success_rate': round(success_rate, 2)
        }

    # ========================================================================
    # CONFIGURATION
    # ========================================================================

    def get_strategies(self) -> dict:
        """Get available strategies"""
        return {
            'strategies': self.strategies,
            'default': 'header_focused'
        }

    def get_default_config(self) -> dict:
        """Get default configuration"""
        return {
            'default_strategy': DEFAULT_CONFIG,
            'homoglyphs_count': len(self.homoglyphs),
            'invisible_chars_count': len(self.invisible_chars)
        }
