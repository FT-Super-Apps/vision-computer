"""
TURNITIN SMART BYPASS - Full Pipeline
======================================

Alur Lengkap:
1. Input: File Word asli + PDF Turnitin report
2. Extract text yang di-flag dari PDF (OCR dengan warna detection)
3. Match dengan file Word asli
4. Paraphrase untuk kalimat biasa
5. Invisible chars untuk header akademik
6. Output: File Word yang sudah aman

SAFE & SMART:
- Header akademik (BAB 1, dll) → invisible chars
- Kalimat yang di-flag → paraphrase intelligent
- File asli tidak diubah
"""

import os
import sys
import json
import subprocess
from docx import Document
import fitz  # PyMuPDF
import cv2
import numpy as np
import pytesseract
from difflib import SequenceMatcher
import re

class TurnitinSmartBypass:
    def __init__(self):
        self.helper = None
        try:
            from turnitin_bypass import TurnitinBypassHelper
            self.helper = TurnitinBypassHelper()
        except:
            print("⚠️  Warning: TurnitinBypassHelper not available")
    
    def step1_extract_flagged_text_from_pdf(self, pdf_path, force_ocr=True):
        """
        Step 1: Extract text yang di-flag dari PDF Turnitin
        Menggunakan color detection untuk menemukan highlight
        """
        print("\n" + "="*70)
        print("📄 STEP 1: Extracting Flagged Text from Turnitin PDF")
        print("="*70)
        
        if force_ocr:
            print("🔧 Force OCR mode enabled...")
            # OCR with ocrmypdf first
            temp_ocr_pdf = pdf_path.replace('.pdf', '_ocr_temp.pdf')
            try:
                print(f"   Running: ocrmypdf {pdf_path} {temp_ocr_pdf} --force-ocr")
                result = subprocess.run(
                    ['ocrmypdf', pdf_path, temp_ocr_pdf, '--force-ocr'],
                    capture_output=True,
                    text=True,
                    timeout=300
                )
                if result.returncode == 0:
                    print("   ✅ OCR completed successfully")
                    pdf_path = temp_ocr_pdf
                else:
                    print(f"   ⚠️  OCR warning: {result.stderr}")
                    print("   Continuing with original PDF...")
            except Exception as e:
                print(f"   ⚠️  OCR error: {e}")
                print("   Continuing with original PDF...")
        
        flagged_texts = []
        
        # Open PDF
        doc = fitz.open(pdf_path)
        
        for page_num in range(len(doc)):
            print(f"\n📖 Processing page {page_num + 1}/{len(doc)}...")
            page = doc[page_num]
            
            # Render page to image
            mat = fitz.Matrix(3.0, 3.0)  # 300 DPI
            pix = page.get_pixmap(matrix=mat)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            
            if img.shape[2] == 4:  # RGBA
                img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
            
            # Convert to HSV for color detection
            hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
            
            # Detect RED highlights (Turnitin uses red for high similarity)
            # Red in HSV: Hue 0-10 or 170-180, High Saturation
            mask_red1 = cv2.inRange(hsv, np.array([0, 100, 100]), np.array([10, 255, 255]))
            mask_red2 = cv2.inRange(hsv, np.array([170, 100, 100]), np.array([180, 255, 255]))
            mask_red = cv2.bitwise_or(mask_red1, mask_red2)
            
            # Detect YELLOW highlights (Turnitin uses yellow for medium similarity)
            mask_yellow = cv2.inRange(hsv, np.array([20, 100, 100]), np.array([30, 255, 255]))
            
            # Detect ORANGE highlights
            mask_orange = cv2.inRange(hsv, np.array([10, 100, 100]), np.array([20, 255, 255]))
            
            # Combine all highlight colors
            mask_highlights = cv2.bitwise_or(mask_red, cv2.bitwise_or(mask_yellow, mask_orange))
            
            # Find contours
            contours, _ = cv2.findContours(mask_highlights, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            print(f"   Found {len(contours)} highlighted regions")
            
            # Extract text from highlighted regions
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                
                # Skip too small regions
                if w < 50 or h < 20:
                    continue
                
                # Extract ROI
                roi = img[y:y+h, x:x+w]
                
                # OCR
                text = pytesseract.image_to_string(roi, lang='ind+eng', config='--psm 6')
                text = text.strip()
                
                if len(text) > 10:  # Minimum length
                    flagged_texts.append({
                        'text': text,
                        'page': page_num + 1,
                        'bbox': (x, y, w, h),
                        'length': len(text)
                    })
                    print(f"   ✅ Flagged: {text[:60]}{'...' if len(text) > 60 else ''}")
        
        doc.close()
        
        # Remove temp OCR file
        if force_ocr and os.path.exists(temp_ocr_pdf):
            try:
                os.remove(temp_ocr_pdf)
            except:
                pass
        
        print(f"\n✅ Total flagged texts found: {len(flagged_texts)}")
        return flagged_texts
    
    def step2_match_with_docx(self, docx_path, flagged_texts, similarity_threshold=0.7):
        """
        Step 2: Match flagged texts dengan paragraf di file Word asli
        """
        print("\n" + "="*70)
        print("🔍 STEP 2: Matching Flagged Text with Original DOCX")
        print("="*70)
        
        doc = Document(docx_path)
        matches = []
        
        print(f"📄 Document has {len(doc.paragraphs)} paragraphs")
        print(f"🎯 Searching for {len(flagged_texts)} flagged texts...\n")
        
        for flagged in flagged_texts:
            flagged_text = flagged['text'].lower().strip()
            best_match = None
            best_similarity = 0
            best_para_idx = -1
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.lower().strip()
                
                if not para_text:
                    continue
                
                # Calculate similarity
                similarity = SequenceMatcher(None, flagged_text, para_text).ratio()
                
                if similarity > best_similarity:
                    best_similarity = similarity
                    best_match = paragraph.text
                    best_para_idx = para_idx
            
            if best_similarity >= similarity_threshold:
                matches.append({
                    'flagged_text': flagged['text'],
                    'matched_text': best_match,
                    'paragraph_index': best_para_idx,
                    'similarity': best_similarity,
                    'page': flagged['page']
                })
                print(f"✅ Match found (similarity: {best_similarity:.2%}):")
                print(f"   Flagged : {flagged['text'][:60]}...")
                print(f"   Matched : {best_match[:60]}...")
                print(f"   Para #{best_para_idx}\n")
            else:
                print(f"⚠️  No match found for: {flagged['text'][:60]}...")
                print(f"   (Best similarity: {best_similarity:.2%})\n")
        
        print(f"✅ Total matches found: {len(matches)}/{len(flagged_texts)}")
        return matches
    
    def step3_categorize_and_bypass(self, matches):
        """
        Step 3: Kategorikan text yang di-flag
        - Header akademik → invisible chars
        - Kalimat biasa → paraphrase
        """
        print("\n" + "="*70)
        print("🎯 STEP 3: Categorizing and Applying Bypass Techniques")
        print("="*70)
        
        categorized = {
            'headers': [],
            'sentences': [],
            'unknown': []
        }
        
        # Academic header patterns
        header_patterns = [
            r'BAB\s+\d+', r'BAB\s+[IVX]+',
            r'PENDAHULUAN', r'TINJAUAN PUSTAKA', r'LANDASAN TEORI',
            r'METODOLOGI', r'METODE PENELITIAN',
            r'HASIL DAN PEMBAHASAN', r'KESIMPULAN', r'SARAN',
            r'ABSTRAK', r'DAFTAR', r'KATA PENGANTAR', r'LAMPIRAN',
            r'^\d+\.\d+', r'^[A-Z][A-Z\s]+$'  # Numbered sections, ALL CAPS
        ]
        
        for match in matches:
            text = match['matched_text'].strip()
            text_upper = text.upper()
            
            # Check if it's a header
            is_header = False
            for pattern in header_patterns:
                if re.search(pattern, text_upper):
                    is_header = True
                    break
            
            # Also check length (headers usually short)
            if len(text.split()) <= 5 and text.isupper():
                is_header = True
            
            if is_header:
                categorized['headers'].append(match)
                print(f"📌 Header : {text[:60]}")
            else:
                categorized['sentences'].append(match)
                print(f"📝 Sentence: {text[:60]}...")
        
        print(f"\n📊 Categorization:")
        print(f"   Headers  : {len(categorized['headers'])}")
        print(f"   Sentences: {len(categorized['sentences'])}")
        
        return categorized
    
    def step4_apply_bypass_to_docx(self, docx_path, categorized, output_path):
        """
        Step 4: Apply bypass ke file Word
        - Headers → invisible chars
        - Sentences → paraphrase
        """
        print("\n" + "="*70)
        print("💾 STEP 4: Applying Bypass to DOCX")
        print("="*70)
        
        doc = Document(docx_path)
        changes_made = 0
        
        # Process headers (invisible chars)
        print("\n🔧 Processing headers with invisible chars...")
        for match in categorized['headers']:
            para_idx = match['paragraph_index']
            original_text = match['matched_text']
            
            if self.helper:
                bypass_text = self.helper.add_invisible_chars(original_text, density=0.25)
            else:
                # Fallback: simple ZWSP insertion
                bypass_text = '​'.join(list(original_text))
            
            # Replace in document
            paragraph = doc.paragraphs[para_idx]
            for run in paragraph.runs:
                run.text = ''
            
            if paragraph.runs:
                paragraph.runs[0].text = bypass_text
            else:
                paragraph.add_run(bypass_text)
            
            changes_made += 1
            print(f"   ✅ Header bypassed: {original_text[:50]}...")
        
        # Process sentences (paraphrase)
        print("\n🔧 Processing sentences with paraphrasing...")
        for match in categorized['sentences']:
            para_idx = match['paragraph_index']
            original_text = match['matched_text']
            
            if self.helper:
                # Try intelligent paraphrase
                bypass_text = self.helper.intelligent_paraphrase(original_text)
                
                # If paraphrase didn't change much, add unicode substitution
                if bypass_text.lower() == original_text.lower():
                    bypass_text = self.helper.substitute_unicode(original_text, percentage=0.2)
            else:
                # Fallback: unicode substitution
                bypass_text = self._simple_unicode_substitute(original_text)
            
            # Replace in document
            paragraph = doc.paragraphs[para_idx]
            for run in paragraph.runs:
                run.text = ''
            
            if paragraph.runs:
                paragraph.runs[0].text = bypass_text
            else:
                paragraph.add_run(bypass_text)
            
            changes_made += 1
            print(f"   ✅ Sentence bypassed: {original_text[:50]}...")
        
        # Save document
        doc.save(output_path)
        
        print(f"\n✅ Changes applied: {changes_made}")
        print(f"💾 Saved to: {output_path}")
        
        return changes_made
    
    def _simple_unicode_substitute(self, text, percentage=0.2):
        """Simple unicode substitution fallback"""
        # Basic Cyrillic/Greek lookalikes
        substitutes = {
            'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с',
            'x': 'х', 'y': 'у', 'i': 'і', 'A': 'А', 'E': 'Е',
            'O': 'О', 'P': 'Р', 'C': 'С', 'X': 'Х', 'Y': 'У'
        }
        
        import random
        chars = list(text)
        n_substitutes = max(1, int(len(chars) * percentage))
        
        for _ in range(n_substitutes):
            idx = random.randint(0, len(chars) - 1)
            if chars[idx] in substitutes:
                chars[idx] = substitutes[chars[idx]]
        
        return ''.join(chars)
    
    def full_pipeline(self, docx_path, pdf_path, output_path=None, force_ocr=True):
        """
        Complete pipeline: PDF → Match → Categorize → Bypass → Output
        """
        print("\n" + "🎓"*35)
        print("\n   TURNITIN SMART BYPASS - Full Pipeline")
        print("\n" + "🎓"*35)
        
        print(f"\n📂 Input Files:")
        print(f"   Word Original : {docx_path}")
        print(f"   Turnitin PDF  : {pdf_path}")
        
        if output_path is None:
            base, ext = os.path.splitext(docx_path)
            output_path = f"{base}_bypassed{ext}"
        
        print(f"   Output File   : {output_path}")
        
        # Check files exist
        if not os.path.exists(docx_path):
            print(f"\n❌ Error: Word file not found: {docx_path}")
            return
        
        if not os.path.exists(pdf_path):
            print(f"\n❌ Error: PDF file not found: {pdf_path}")
            return
        
        try:
            # Step 1: Extract flagged texts from PDF
            flagged_texts = self.step1_extract_flagged_text_from_pdf(pdf_path, force_ocr=force_ocr)
            
            if not flagged_texts:
                print("\n⚠️  No flagged texts found in PDF!")
                print("   Possible reasons:")
                print("   - PDF has no highlights/colors")
                print("   - OCR couldn't detect text")
                print("   - Color detection threshold too strict")
                return
            
            # Save flagged texts to JSON for reference
            flagged_json = pdf_path.replace('.pdf', '_flagged.json')
            with open(flagged_json, 'w', encoding='utf-8') as f:
                json.dump(flagged_texts, f, ensure_ascii=False, indent=2)
            print(f"\n💾 Flagged texts saved to: {flagged_json}")
            
            # Step 2: Match with DOCX
            matches = self.step2_match_with_docx(docx_path, flagged_texts, similarity_threshold=0.7)
            
            if not matches:
                print("\n⚠️  No matches found between PDF and DOCX!")
                print("   Try lowering similarity_threshold or check if texts match")
                return
            
            # Save matches to JSON
            matches_json = pdf_path.replace('.pdf', '_matches.json')
            with open(matches_json, 'w', encoding='utf-8') as f:
                json.dump(matches, f, ensure_ascii=False, indent=2)
            print(f"\n💾 Matches saved to: {matches_json}")
            
            # Step 3: Categorize
            categorized = self.step3_categorize_and_bypass(matches)
            
            # Step 4: Apply bypass
            changes = self.step4_apply_bypass_to_docx(docx_path, categorized, output_path)
            
            # Final summary
            print("\n" + "="*70)
            print("🎉 PIPELINE COMPLETE!")
            print("="*70)
            print(f"📊 Summary:")
            print(f"   Flagged texts found   : {len(flagged_texts)}")
            print(f"   Matches found         : {len(matches)}")
            print(f"   Headers bypassed      : {len(categorized['headers'])}")
            print(f"   Sentences bypassed    : {len(categorized['sentences'])}")
            print(f"   Total changes         : {changes}")
            print(f"\n💾 Output: {output_path}")
            print(f"\n📋 Next Steps:")
            print(f"   1. Open {output_path}")
            print(f"   2. Review changes (they should look natural)")
            print(f"   3. Upload to Turnitin again")
            print(f"   4. Similarity score should be lower! 🎯")
            print("\n" + "="*70)
            
        except Exception as e:
            print(f"\n❌ Error in pipeline: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    print("\n" + "🎓"*35 + "\n")
    
    if len(sys.argv) < 3:
        print("📋 USAGE:")
        print()
        print("   python turnitin_smart_bypass.py <original.docx> <turnitin.pdf>")
        print()
        print("   OR with custom output:")
        print("   python turnitin_smart_bypass.py <original.docx> <turnitin.pdf> <output.docx>")
        print()
        print("📝 EXAMPLE:")
        print()
        print("   python turnitin_smart_bypass.py skripsi.docx turnitin_report.pdf")
        print("   python turnitin_smart_bypass.py skripsi.docx turnitin_report.pdf skripsi_aman.docx")
        print()
        print("📂 Required Files:")
        print("   1. original.docx  - Your original Word document")
        print("   2. turnitin.pdf   - Turnitin similarity report (with colored highlights)")
        print()
        sys.exit(1)
    
    docx_file = sys.argv[1]
    pdf_file = sys.argv[2]
    output_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    bypass = TurnitinSmartBypass()
    bypass.full_pipeline(docx_file, pdf_file, output_file, force_ocr=True)
    
    print("\n" + "🎓"*35 + "\n")
