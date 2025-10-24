#!/usr/bin/env python3
"""
Smart Content Analyzer
Analyzes document to detect potentially plagiarized content
without needing external flag.txt file
"""

import re
from difflib import SequenceMatcher
from pathlib import Path
from docx import Document
from typing import List, Dict, Tuple
import statistics

class ContentAnalyzer:
    """Analyzes document content for plagiarism risk"""

    def __init__(self):
        # Patterns untuk detect content yang likely ter-flag
        self.common_academic_starters = [
            'menurut',
            'berdasarkan',
            'penelitian',
            'peneliti',
            'hasil',
            'analisis',
            'kesimpulan',
            'temuan',
            'observasi',
            'pengamatan',
            'studi',
            'dengan demikian',
            'oleh karena itu',
            'sebab itu',
            'maka dari itu'
        ]

        self.direct_quote_indicators = [
            '"',
            '"',
            '"',
            '„',
            '‟',
            '"',
            '"'
        ]

    # ========================================================================
    # CONTENT QUALITY METRICS
    # ========================================================================

    def calculate_sentence_length(self, text: str) -> Dict[str, float]:
        """Calculate average sentence length"""
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 10]

        if not sentences:
            return {'min': 0, 'max': 0, 'avg': 0, 'count': 0}

        lengths = [len(s.split()) for s in sentences]
        return {
            'min': min(lengths) if lengths else 0,
            'max': max(lengths) if lengths else 0,
            'avg': statistics.mean(lengths) if lengths else 0,
            'count': len(sentences)
        }

    def detect_quote_blocks(self, text: str) -> List[Dict]:
        """Detect quoted sections (high plagiarism risk)"""
        quotes = []

        # Find quoted text
        pattern = r'["\"\"\"\„\‟\"\"]([^"]*?)["\"\"\"\„\‟\"\"]'
        matches = re.finditer(pattern, text)

        for match in matches:
            quote_text = match.group(1).strip()
            if len(quote_text) > 20:  # Only significant quotes
                quotes.append({
                    'text': quote_text[:100],
                    'length': len(quote_text),
                    'risk': 'very_high'  # Direct quotes always high risk
                })

        return quotes

    def detect_unusual_vocabulary(self, paragraph: str) -> List[str]:
        """Detect words that don't match surrounding context (copy-paste indicator)"""
        # Simple heuristic: words with unusual character patterns
        words = paragraph.split()

        suspicious = []
        for word in words:
            # Check for mixed case, special patterns
            if re.match(r'^[A-Z][a-z]+[A-Z]', word):  # CamelCase
                suspicious.append(word)
            elif re.match(r'^[a-z]+\d+[a-z]*', word):  # Numbers in word
                suspicious.append(word)
            elif re.match(r'^[A-Z]{2,}', word) and len(word) > 3:  # Acronyms
                suspicious.append(word)

        return suspicious[:5]  # Top 5

    def measure_lexical_diversity(self, text: str) -> float:
        """
        Measure uniqueness of vocabulary
        Lower = more repetition = higher plagiarism risk
        """
        words = text.lower().split()
        if not words:
            return 0

        unique_words = len(set(words))
        total_words = len(words)

        # Type-Token Ratio (TTR)
        diversity = unique_words / total_words if total_words > 0 else 0
        return diversity

    def detect_formatting_inconsistencies(self, doc: Document) -> List[Dict]:
        """Detect formatting changes (copy-paste indicator)"""
        inconsistencies = []

        font_changes = {}
        for i, para in enumerate(doc.paragraphs):
            if not para.runs:
                continue

            for run in para.runs:
                if run.font.name:
                    key = f"{i}:{run.font.name}:{run.font.size}"
                    font_changes[key] = font_changes.get(key, 0) + 1

        # Detect sudden font changes
        if len(font_changes) > len(doc.paragraphs) // 2:
            inconsistencies.append({
                'type': 'font_changes',
                'count': len(font_changes),
                'risk': 'medium',
                'detail': 'Banyak perubahan font dalam dokumen'
            })

        return inconsistencies

    # ========================================================================
    # PLAGIARISM RISK DETECTION
    # ========================================================================

    def detect_potential_plagiarism_paragraphs(self, doc: Document) -> List[Dict]:
        """
        Detect paragraphs with high plagiarism risk based on:
        1. Similarity to academic templates
        2. Unusual vocabulary
        3. Quote indicators
        4. Lexical diversity
        """
        risky_paragraphs = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if len(text) < 30:
                continue

            risk_score = 0.0
            risk_factors = []

            # 1. Check for quotes
            if self.detect_quote_blocks(text):
                risk_score += 0.3
                risk_factors.append('quoted_text')

            # 2. Check for direct academic phrasing
            lower_text = text.lower()
            for starter in self.common_academic_starters:
                if lower_text.startswith(starter):
                    risk_score += 0.1
                    risk_factors.append('academic_pattern')
                    break

            # 3. Check unusual vocabulary
            suspicious_words = self.detect_unusual_vocabulary(text)
            if suspicious_words:
                risk_score += 0.15
                risk_factors.append(f'unusual_vocab:{len(suspicious_words)}')

            # 4. Check lexical diversity (very low = repetitive = plagarized)
            diversity = self.measure_lexical_diversity(text)
            if diversity < 0.4:  # Low diversity indicator
                risk_score += 0.2
                risk_factors.append(f'low_diversity:{diversity:.2f}')

            # 5. Check if paragraph is too long (template fill-in indicator)
            word_count = len(text.split())
            if word_count > 200:
                risk_score += 0.1
                risk_factors.append(f'long_para:{word_count}')

            if risk_score > 0.3:  # Threshold
                risky_paragraphs.append({
                    'para_num': i + 1,
                    'text': text[:150],
                    'risk_score': min(risk_score, 1.0),  # Cap at 1.0
                    'risk_level': self._score_to_level(risk_score),
                    'word_count': word_count,
                    'diversity': round(diversity, 3),
                    'factors': risk_factors
                })

        return risky_paragraphs

    def detect_header_content(self, doc: Document) -> List[Dict]:
        """Extract and identify headers/structure"""
        headers = []

        for para in doc.paragraphs:
            text = para.text.strip()

            # Header detection
            if re.match(r'^(BAB|CHAPTER)\s+[IVX]+', text, re.IGNORECASE):
                headers.append({'type': 'chapter', 'text': text})
            elif re.match(r'^[A-Z]\.\s+', text):
                headers.append({'type': 'section', 'text': text})
            elif re.match(r'^[0-9]+\.\s+', text):
                headers.append({'type': 'numbered', 'text': text})
            elif re.match(r'^(PENDAHULUAN|LATAR BELAKANG|TINJAUAN PUSTAKA|METODOLOGI|HASIL|KESIMPULAN)', text, re.IGNORECASE):
                headers.append({'type': 'standard', 'text': text})

        return headers

    # ========================================================================
    # OVERALL ANALYSIS
    # ========================================================================

    def analyze_full_document(self, input_path: str) -> Dict:
        """
        Comprehensive document analysis for plagiarism risk
        """
        doc = Document(input_path)

        # Basic metrics
        total_paragraphs = len(doc.paragraphs)
        total_text = '\n'.join([p.text for p in doc.paragraphs])
        total_words = len(total_text.split())
        total_chars = len(total_text)

        # Analyze
        risky_paragraphs = self.detect_potential_plagiarism_paragraphs(doc)
        headers = self.detect_header_content(doc)
        formatting_issues = self.detect_formatting_inconsistencies(doc)
        sentence_metrics = self.calculate_sentence_length(total_text)
        overall_diversity = self.measure_lexical_diversity(total_text)

        # Calculate overall risk score
        plagiarism_risk = self._calculate_overall_risk(
            len(risky_paragraphs),
            total_paragraphs,
            overall_diversity,
            len(headers)
        )

        return {
            'success': True,
            'filename': Path(input_path).name,
            'document_metrics': {
                'total_paragraphs': total_paragraphs,
                'total_words': total_words,
                'total_characters': total_chars,
                'average_paragraph_length': total_words // total_paragraphs if total_paragraphs > 0 else 0,
                'lexical_diversity': round(overall_diversity, 3)
            },
            'sentence_metrics': sentence_metrics,
            'structure': {
                'headers_found': len(headers),
                'header_types': self._count_header_types(headers)
            },
            'plagiarism_risk': {
                'overall_score': round(plagiarism_risk, 3),
                'estimated_similarity': self._risk_to_similarity(plagiarism_risk),
                'risky_paragraphs_count': len(risky_paragraphs),
                'risk_percentage': round((len(risky_paragraphs) / total_paragraphs * 100) if total_paragraphs > 0 else 0, 1)
            },
            'flagged_content': {
                'high_risk_paragraphs': risky_paragraphs[:20],  # Top 20
                'formatting_issues': formatting_issues,
                'potentially_quoted': self._find_quoted_paragraphs(doc)
            },
            'recommendations': self._generate_recommendations(plagiarism_risk, len(risky_paragraphs), total_paragraphs),
            'analysis_summary': {
                'total_flags_detected': len(risky_paragraphs),
                'critical_sections': sum(1 for p in risky_paragraphs if p['risk_score'] > 0.7),
                'medium_sections': sum(1 for p in risky_paragraphs if 0.4 < p['risk_score'] <= 0.7),
                'low_sections': sum(1 for p in risky_paragraphs if 0.3 < p['risk_score'] <= 0.4)
            }
        }

    # ========================================================================
    # HELPER FUNCTIONS
    # ========================================================================

    def _score_to_level(self, score: float) -> str:
        """Convert risk score to level"""
        if score < 0.4:
            return 'low'
        elif score < 0.6:
            return 'medium'
        elif score < 0.8:
            return 'high'
        else:
            return 'critical'

    def _calculate_overall_risk(self, risky_count: int, total_para: int, diversity: float, headers: int) -> float:
        """Calculate overall plagiarism risk score (0-1)"""
        if total_para == 0:
            return 0

        # Risk from plagiarised paragraphs
        risk_from_paras = (risky_count / total_para) * 0.6

        # Risk from low diversity
        diversity_risk = (1 - diversity) * 0.3 if diversity < 0.5 else 0

        # Risk from missing structure
        structure_risk = 0.1 if headers < 3 else 0

        total_risk = risk_from_paras + diversity_risk + structure_risk
        return min(total_risk, 1.0)

    def _risk_to_similarity(self, risk: float) -> str:
        """Convert risk score to estimated similarity percentage"""
        if risk < 0.3:
            return '10-20%'
        elif risk < 0.5:
            return '20-35%'
        elif risk < 0.7:
            return '35-50%'
        else:
            return '>50%'

    def _count_header_types(self, headers: List[Dict]) -> Dict:
        """Count header types"""
        counts = {}
        for h in headers:
            counts[h['type']] = counts.get(h['type'], 0) + 1
        return counts

    def _find_quoted_paragraphs(self, doc: Document) -> List[str]:
        """Find paragraphs containing quotes"""
        quoted = []
        for para in doc.paragraphs:
            if self.detect_quote_blocks(para.text):
                quoted.append(para.text[:100])
        return quoted

    def _generate_recommendations(self, risk: float, risky_count: int, total: int) -> List[str]:
        """Generate recommendations based on analysis"""
        recommendations = []

        if risk > 0.7:
            recommendations.append('⚠️ CRITICAL: Document has very high plagiarism risk (>70%)')
            recommendations.append('Gunakan Header-Focused bypass dengan density 95% + 40%')

        if risk > 0.5:
            recommendations.append('⚠️ HIGH: Banyak section dengan plagiarism risk')
            recommendations.append('Focus pada top risky paragraphs terlebih dahulu')

        if risky_count > (total * 0.3):
            recommendations.append('📊: >30% paragraphs memiliki plagiarism risk')
            recommendations.append('Pertimbangkan untuk rewrite beberapa section')

        if risk < 0.3:
            recommendations.append('✅ GOOD: Document has low plagiarism risk')
            recommendations.append('Gunakan Natural strategy (light bypass)')

        return recommendations


# Quick test
if __name__ == '__main__':
    analyzer = ContentAnalyzer()

    # Test dengan original.docx jika ada
    test_file = 'original.docx'
    if Path(test_file).exists():
        result = analyzer.analyze_full_document(test_file)
        import json
        print(json.dumps(result, indent=2, ensure_ascii=False))
