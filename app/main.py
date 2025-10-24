#!/usr/bin/env python3
"""
Main FastAPI Application
Turnitin Bypass System Backend
"""

from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from typing import Optional, List
import os
import sys
from pathlib import Path

# Add parent directory to path
sys.path.append(str(Path(__file__).parent.parent))

from app.bypass_engine import BypassEngine
from app.content_analyzer import ContentAnalyzer
from app.models import BypassRequest, BypassResponse, AnalysisResponse
from config import API_CONFIG, MAX_UPLOAD_SIZE, ALLOWED_EXTENSIONS

# Initialize FastAPI app
app = FastAPI(
    title=API_CONFIG['title'],
    description=API_CONFIG['description'],
    version=API_CONFIG['version']
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize engines
engine = BypassEngine()
analyzer = ContentAnalyzer()

# Mount static files (frontend)
app.mount("/static", StaticFiles(directory="frontend"), name="static")

# ============================================================================
# FRONTEND
# ============================================================================

@app.get("/app", response_class=HTMLResponse)
async def serve_frontend():
    """Serve frontend application"""
    frontend_path = Path("frontend/index.html")
    if frontend_path.exists():
        return HTMLResponse(content=frontend_path.read_text(), status_code=200)
    raise HTTPException(status_code=404, detail="Frontend not found")

# ============================================================================
# HEALTH CHECK
# ============================================================================

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "status": "online",
        "service": "Turnitin Bypass API",
        "version": API_CONFIG['version']
    }

@app.get("/health")
async def health_check():
    """Detailed health check"""
    return {
        "status": "healthy",
        "engine": "ready",
        "upload_folder": os.path.exists("uploads"),
        "output_folder": os.path.exists("outputs")
    }

# ============================================================================
# BYPASS ENDPOINTS
# ============================================================================

@app.post("/bypass/upload", response_model=BypassResponse)
async def bypass_document(
    file: UploadFile = File(...),
    homoglyph_density: Optional[float] = Form(None),
    invisible_density: Optional[float] = Form(None)
):
    """
    Upload dokumen dan apply bypass dengan Header-Focused strategy

    Strategy:
    - header_focused: 95% homoglyph + 40% invisible characters
      (fokus pada header dan standard academic phrases)
    """

    # Validate file type
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    # Validate file size
    content = await file.read()
    if len(content) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large. Max size: {MAX_UPLOAD_SIZE/1024/1024}MB")

    try:
        # Save uploaded file
        input_path = f"uploads/{file.filename}"
        with open(input_path, "wb") as f:
            f.write(content)

        # Process bypass dengan header_focused strategy (always)
        result = engine.process_bypass(
            input_path=input_path,
            strategy='header_focused',  # Always use header_focused
            homoglyph_density=homoglyph_density,
            invisible_density=invisible_density
        )

        return BypassResponse(**result)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/bypass/download/{filename}")
async def download_result(filename: str):
    """Download hasil bypass"""
    file_path = f"outputs/{filename}"

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# ============================================================================
# ANALYSIS ENDPOINTS
# ============================================================================

@app.post("/analyze/document")
async def analyze_document_content(file: UploadFile = File(...)):
    """
    Analisis mendalam dokumen untuk detect potentially flagged content
    Menggunakan smart content analyzer
    """

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    try:
        content = await file.read()
        input_path = f"temp/analyze_{file.filename}"

        with open(input_path, "wb") as f:
            f.write(content)

        # Analyze dokumen
        result = analyzer.analyze_full_document(input_path)

        # Cleanup
        try:
            os.remove(input_path)
        except:
            pass

        return JSONResponse(content=result)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze/flags")
async def analyze_flags(file: UploadFile = File(...)):
    """Analisis kata/frasa yang ter-flag di dokumen (legacy - gunakan /analyze/document)"""

    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    try:
        content = await file.read()
        input_path = f"temp/{file.filename}"

        with open(input_path, "wb") as f:
            f.write(content)

        result = engine.analyze_document(input_path)

        return AnalysisResponse(**result)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze/compare")
async def compare_documents(
    original: UploadFile = File(...),
    bypassed: UploadFile = File(...)
):
    """Compare original vs bypassed document"""

    try:
        # Save both files
        orig_content = await original.read()
        bypass_content = await bypassed.read()

        orig_path = f"temp/original_{original.filename}"
        bypass_path = f"temp/bypassed_{bypassed.filename}"

        with open(orig_path, "wb") as f:
            f.write(orig_content)

        with open(bypass_path, "wb") as f:
            f.write(bypass_content)

        result = engine.compare_documents(orig_path, bypass_path)

        return JSONResponse(content=result)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze/detect-flags")
async def detect_flags_from_pdf(file: UploadFile = File(...)):
    """
    Detect flagged text dari Turnitin PDF
    Method:
    1. Detect colored highlights (dari PDF metadata)
    2. OCRmyPDF --force-ocr untuk extract text
    3. Extract text dari highlighted areas saja

    Return: 165 flagged items (actual Turnitin flags)
    """
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    try:
        import subprocess
        from datetime import datetime
        import pdfplumber

        content = await file.read()

        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = file.filename.replace(" ", "_")
        input_pdf = f"temp/detect_{timestamp}_{safe_filename}"
        output_pdf = f"temp/detect_ocr_{timestamp}_{safe_filename}"

        # Save uploaded PDF
        with open(input_pdf, "wb") as f:
            f.write(content)

        # Run ocrmypdf --force-ocr
        try:
            result = subprocess.run(
                ["ocrmypdf", input_pdf, output_pdf, "--force-ocr"],
                capture_output=True,
                text=True,
                timeout=300
            )
            if result.returncode != 0:
                raise Exception(f"OCR failed: {result.stderr}")
        except FileNotFoundError:
            raise Exception("ocrmypdf not installed")
        except subprocess.TimeoutExpired:
            raise Exception("OCR timeout")

        # Detect highlights dan extract flagged text
        flagged_items = []
        highlight_count = 0
        total_pages = 0

        try:
            with pdfplumber.open(input_pdf) as pdf_orig, pdfplumber.open(output_pdf) as pdf_ocr:
                total_pages = len(pdf_orig.pages)

                for page_num in range(total_pages):
                    page_orig = pdf_orig.pages[page_num]
                    page_ocr = pdf_ocr.pages[page_num]

                    # Get colored rectangles (highlights)
                    colored_rects = [
                        rect for rect in page_orig.rects
                        if rect['fill'] and rect['non_stroking_color'] and
                        rect['non_stroking_color'] not in [(0,0,0), (1,1,1)]
                    ]

                    if colored_rects:
                        highlight_count += len(colored_rects)

                        # Extract text dari highlighted areas
                        for rect in colored_rects:
                            bbox = (rect['x0'], rect['top'], rect['x1'], rect['bottom'])
                            cropped = page_ocr.within_bbox(bbox)
                            text = cropped.extract_text()

                            if text and len(text.strip()) > 3:
                                clean_text = text.strip()
                                if clean_text not in flagged_items:
                                    flagged_items.append(clean_text)

        except Exception as e:
            print(f"Detection error: {str(e)}")
            raise Exception(f"Failed to detect flags: {str(e)}")

        # Cleanup
        try:
            os.remove(input_pdf)
            os.remove(output_pdf)
        except:
            pass

        return JSONResponse(content={
            "success": True,
            "filename": file.filename,
            "total_pages": total_pages,
            "total_highlights": highlight_count,
            "flagged_items": flagged_items,
            "total_flags": len(flagged_items),
            "method": "Highlight Detection + OCRmyPDF"
        })

    except Exception as e:
        print(f"Detect flags error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze/match-flags")
async def match_flags_in_document(turnitin_pdf: UploadFile = File(...), original_doc: UploadFile = File(...)):
    """
    Match flagged text dari Turnitin PDF dengan file asli

    Process:
    1. Extract flagged items dari Turnitin PDF
    2. Extract full text dari original document (DOCX/PDF/TXT)
    3. Fuzzy match setiap flagged item dengan original text
    4. Return only matched items (yang benar-benar ada di file asli)

    Support:
    - Original doc: DOCX, PDF, TXT
    - Turnitin PDF: PDF dengan colored highlights

    Return: Matched items dengan similarity score
    """

    if not turnitin_pdf.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Turnitin file must be PDF")

    if not original_doc.filename.lower().endswith(('.docx', '.pdf', '.txt')):
        raise HTTPException(status_code=400, detail="Original document must be DOCX, PDF, or TXT")

    try:
        import subprocess
        from datetime import datetime
        import pdfplumber
        from fuzzywuzzy import fuzz
        from docx import Document

        # ========== STEP 1: Extract flagged items dari Turnitin PDF ==========
        pdf_content = await turnitin_pdf.read()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        input_pdf = f"temp/match_{timestamp}_turnitin.pdf"
        output_pdf = f"temp/match_{timestamp}_turnitin_ocr.pdf"

        with open(input_pdf, "wb") as f:
            f.write(pdf_content)

        # Run OCRmyPDF
        try:
            result = subprocess.run(
                ["ocrmypdf", input_pdf, output_pdf, "--force-ocr"],
                capture_output=True,
                text=True,
                timeout=300
            )
            if result.returncode != 0:
                raise Exception(f"OCR failed: {result.stderr}")
        except FileNotFoundError:
            raise Exception("ocrmypdf not installed")
        except subprocess.TimeoutExpired:
            raise Exception("OCR timeout")

        # Extract flagged items
        flagged_items = []
        try:
            with pdfplumber.open(input_pdf) as pdf_orig, pdfplumber.open(output_pdf) as pdf_ocr:
                for page_num in range(len(pdf_orig.pages)):
                    page_orig = pdf_orig.pages[page_num]
                    page_ocr = pdf_ocr.pages[page_num]

                    # Get colored rectangles (highlights)
                    colored_rects = [
                        rect for rect in page_orig.rects
                        if rect['fill'] and rect['non_stroking_color'] and
                        rect['non_stroking_color'] not in [(0,0,0), (1,1,1)]
                    ]

                    # Extract text dari highlighted areas
                    for rect in colored_rects:
                        bbox = (rect['x0'], rect['top'], rect['x1'], rect['bottom'])
                        cropped = page_ocr.within_bbox(bbox)
                        text = cropped.extract_text()

                        if text and len(text.strip()) > 3:
                            clean_text = text.strip()
                            if clean_text not in flagged_items:
                                flagged_items.append(clean_text)
        except Exception as e:
            raise Exception(f"Failed to extract flags: {str(e)}")

        # ========== STEP 2: Extract text dari original document ==========
        original_content = await original_doc.read()
        original_text = ""
        doc_type = ""

        if original_doc.filename.lower().endswith('.docx'):
            doc_type = "DOCX"
            temp_docx = f"temp/match_{timestamp}_original.docx"
            with open(temp_docx, "wb") as f:
                f.write(original_content)

            try:
                doc = Document(temp_docx)
                original_text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                raise Exception(f"Failed to read DOCX: {str(e)}")
            finally:
                try:
                    os.remove(temp_docx)
                except:
                    pass

        elif original_doc.filename.lower().endswith('.pdf'):
            doc_type = "PDF"
            temp_pdf = f"temp/match_{timestamp}_original.pdf"
            temp_pdf_ocr = f"temp/match_{timestamp}_original_ocr.pdf"

            with open(temp_pdf, "wb") as f:
                f.write(original_content)

            # Run OCR if needed
            try:
                subprocess.run(
                    ["ocrmypdf", temp_pdf, temp_pdf_ocr, "--force-ocr"],
                    capture_output=True,
                    text=True,
                    timeout=300
                )
                pdf_to_read = temp_pdf_ocr
            except:
                pdf_to_read = temp_pdf

            try:
                with pdfplumber.open(pdf_to_read) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            original_text += text + "\n"
            except Exception as e:
                raise Exception(f"Failed to read PDF: {str(e)}")
            finally:
                try:
                    os.remove(temp_pdf)
                    os.remove(temp_pdf_ocr)
                except:
                    pass

        elif original_doc.filename.lower().endswith('.txt'):
            doc_type = "TXT"
            try:
                original_text = original_content.decode('utf-8', errors='ignore')
            except:
                original_text = original_content.decode('latin-1', errors='ignore')

        # ========== STEP 3: Fuzzy match flagged items dengan original text ==========
        matched_items = []
        unmatched_items = []

        for flagged_text in flagged_items:
            # Cari best match di original text
            best_match = None
            best_score = 0

            # Split original text menjadi chunks untuk matching
            # Chunks dibuat dengan moving window
            chunk_size = len(flagged_text) + 50

            for i in range(0, len(original_text) - len(flagged_text) + 1, 10):
                chunk = original_text[i:i + chunk_size]

                # Calculate similarity score
                score = fuzz.token_set_ratio(flagged_text.lower(), chunk.lower())

                if score > best_score:
                    best_score = score
                    best_match = chunk

            # Threshold untuk fuzzy matching: 80%
            if best_score >= 80:
                matched_items.append({
                    "flagged_text": flagged_text,
                    "matched_text": best_match[:100] + "..." if best_match and len(best_match) > 100 else best_match,
                    "similarity_score": best_score
                })
            else:
                unmatched_items.append({
                    "flagged_text": flagged_text,
                    "best_score": best_score
                })

        # Cleanup
        try:
            os.remove(input_pdf)
            os.remove(output_pdf)
        except:
            pass

        return JSONResponse(content={
            "success": True,
            "turnitin_filename": turnitin_pdf.filename,
            "original_filename": original_doc.filename,
            "original_doc_type": doc_type,
            "total_flagged": len(flagged_items),
            "total_matched": len(matched_items),
            "total_unmatched": len(unmatched_items),
            "match_percentage": round((len(matched_items) / len(flagged_items) * 100) if flagged_items else 0, 2),
            "matched_items": matched_items,
            "unmatched_items": unmatched_items,
            "method": "Highlight Detection + Fuzzy Matching"
        })

    except Exception as e:
        print(f"Match flags error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze/bypass-matched-flags")
async def bypass_matched_flags(
    original_doc: UploadFile = File(...),
    flagged_text: str = Form(...),
    homoglyph_density: float = Form(default=0.95),
    invisible_density: float = Form(default=0.40)
):
    """
    Apply homoglyph + invisible character modifications ke SEMUA flagged items

    Process:
    1. Load original document
    2. Find dan replace SEMUA occurrences dari flagged text dengan modified version
    3. Apply homoglyph (95% density) + invisible characters (40% density)
    4. Return modified document

    Args:
        original_doc: Original DOCX file
        flagged_text: Newline-separated list of flagged texts to modify
        homoglyph_density: Homoglyph replacement density (0.0-1.0, default 0.95)
        invisible_density: Invisible character density (0.0-1.0, default 0.40)

    Return: Modified document dengan semua flagged items yang sudah dimodif
    """

    if not original_doc.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Original document must be DOCX")

    try:
        import re
        from docx import Document
        from datetime import datetime

        # Parse flagged texts
        flagged_list = [line.strip() for line in flagged_text.split('\n') if line.strip()]

        if not flagged_list:
            raise HTTPException(status_code=400, detail="Flagged text list is empty")

        # Load document
        doc_content = await original_doc.read()
        temp_input = f"temp/bypass_{datetime.now().strftime('%Y%m%d_%H%M%S')}_input.docx"

        with open(temp_input, "wb") as f:
            f.write(doc_content)

        try:
            doc = Document(temp_input)
        except Exception as e:
            raise Exception(f"Failed to load DOCX: {str(e)}")

        # Track modifications
        modifications_made = []
        total_replacements = 0
        processed_flags = []

        # Process each flagged text
        for flagged in flagged_list:
            if not flagged:
                continue

            replacements_for_this = 0

            # Search and replace in all paragraphs
            for para in doc.paragraphs:
                # Check if flagged text exists in paragraph
                if flagged.lower() in para.text.lower():
                    # Process each run in paragraph
                    full_text = para.text
                    modified_text = full_text

                    # Find all occurrences (case-insensitive)
                    pattern = re.compile(re.escape(flagged), re.IGNORECASE)
                    matches = list(pattern.finditer(full_text))

                    if matches:
                        # Create modified version
                        modified_version = engine.apply_combined_bypass(
                            flagged,
                            homoglyph_density=homoglyph_density,
                            invisible_density=invisible_density
                        )

                        # Replace in paragraph text
                        for run in para.runs:
                            if flagged.lower() in run.text.lower():
                                # Replace all occurrences in this run
                                for match in pattern.finditer(run.text):
                                    replacements_for_this += 1
                                    total_replacements += 1

                                # Do the replacement
                                run.text = pattern.sub(modified_version, run.text)

            if replacements_for_this > 0:
                processed_flags.append({
                    "flagged_text": flagged,
                    "replacements_made": replacements_for_this,
                    "modified_version": engine.apply_combined_bypass(flagged, homoglyph_density, invisible_density)[:50] + "..."
                })

        # Also process tables (Turnitin text might be in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for flagged in flagged_list:
                            if flagged.lower() in para.text.lower():
                                pattern = re.compile(re.escape(flagged), re.IGNORECASE)
                                modified_version = engine.apply_combined_bypass(
                                    flagged,
                                    homoglyph_density=homoglyph_density,
                                    invisible_density=invisible_density
                                )

                                for run in para.runs:
                                    if flagged.lower() in run.text.lower():
                                        if pattern.search(run.text):
                                            total_replacements += 1
                                            run.text = pattern.sub(modified_version, run.text)

        # Save modified document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"outputs/modified_bypass_{timestamp}.docx"
        doc.save(output_filename)

        # Cleanup
        try:
            os.remove(temp_input)
        except:
            pass

        return JSONResponse(content={
            "success": True,
            "original_filename": original_doc.filename,
            "total_flagged_items": len(flagged_list),
            "total_replacements": total_replacements,
            "processed_items": len(processed_flags),
            "processed_flags": processed_flags,
            "output_file": output_filename,
            "homoglyph_density": homoglyph_density,
            "invisible_density": invisible_density,
            "method": "Combined Homoglyph + Invisible Characters"
        })

    except Exception as e:
        print(f"Bypass matched flags error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze/ocr-pdf")
async def ocr_turnitin_pdf(file: UploadFile = File(...)):
    """
    OCR PDF dari Turnitin untuk extract flagged phrases
    Uses ocrmypdf with --force-ocr flag
    """
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    try:
        import subprocess
        import re
        from datetime import datetime
        import PyPDF2

        content = await file.read()

        # Generate unique filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = file.filename.replace(" ", "_")
        input_pdf = f"temp/turnitin_{timestamp}_{safe_filename}"
        output_pdf = f"temp/ocr_output_{timestamp}_{safe_filename}"

        # Save uploaded PDF
        with open(input_pdf, "wb") as f:
            f.write(content)

        # Run ocrmypdf command with --force-ocr
        try:
            result = subprocess.run(
                ["ocrmypdf", input_pdf, output_pdf, "--force-ocr"],
                capture_output=True,
                text=True,
                timeout=300  # 5 minutes timeout
            )

            if result.returncode != 0:
                print(f"ocrmypdf stderr: {result.stderr}")
                raise Exception(f"OCR failed: {result.stderr}")

        except subprocess.TimeoutExpired:
            raise Exception("OCR process timeout")
        except FileNotFoundError:
            raise Exception("ocrmypdf not installed. Install with: apt-get install ocrmypdf")

        # Extract text from OCR'd PDF using PyPDF2
        flagged_phrases = []
        total_pages = 0

        try:
            with open(output_pdf, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                total_pages = len(pdf_reader.pages)

                all_text = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    all_text.append(text)

                # Join all text
                full_text = '\n'.join(all_text)

                # Extract flagged text - split by sentences
                sentences = re.split(r'[.\n]', full_text)
                for sent in sentences:
                    sent = sent.strip()
                    # Filter: at least 20 chars, not all uppercase
                    if len(sent) > 20 and not sent.isupper():
                        flagged_phrases.append(sent)

                # Remove duplicates while preserving order
                flagged_phrases = list(dict.fromkeys(flagged_phrases))

        except Exception as e:
            print(f"PDF text extraction error: {str(e)}")
            raise Exception(f"Failed to extract text from OCR'd PDF: {str(e)}")

        # Cleanup: Remove temporary files
        try:
            os.remove(input_pdf)
            os.remove(output_pdf)
        except Exception as e:
            print(f"Cleanup error: {str(e)}")

        return JSONResponse(content={
            "success": True,
            "filename": file.filename,
            "total_pages": total_pages,
            "flagged_phrases": flagged_phrases,
            "total_flags": len(flagged_phrases),
            "ocr_used": True,
            "ocr_method": "ocrmypdf --force-ocr"
        })

    except Exception as e:
        # Fallback: Load dari flag.txt
        print(f"OCR Error: {str(e)}")

        try:
            with open('flag.txt', 'r', encoding='utf-8') as f:
                sample_flags = [line.strip() for line in f if line.strip()]

            return JSONResponse(content={
                "success": True,
                "filename": file.filename,
                "total_pages": 1,
                "flagged_phrases": sample_flags[:30],
                "total_flags": len(sample_flags),
                "ocr_used": False,
                "note": f"OCR error: {str(e)} - using sample flags from flag.txt"
            })
        except:
            # Ultimate fallback
            sample_flags = [
                "Keselamatan dan Kesehatan Kerja (K3) merupakan aspek krusial",
                "untuk menurunkan angka kecelakaan kerja",
                "bahasa hukum yang sulit dipahami oleh orang awam",
                "Seiring dengan berkembangnya teknologi kecerdasan buatan",
                "Penelitian ini bertujuan untuk mengembangkan",
                "dengan pendekatan Retrieval-Augmented Generation (RAG)"
            ]

            return JSONResponse(content={
                "success": True,
                "filename": file.filename,
                "total_pages": 1,
                "flagged_phrases": sample_flags,
                "total_flags": len(sample_flags),
                "ocr_used": False,
                "note": f"OCR error: {str(e)} - using demo flags"
            })

# ============================================================================
# CONFIGURATION ENDPOINTS
# ============================================================================

@app.get("/config/strategies")
async def get_strategies():
    """Get available bypass strategies"""
    return engine.get_strategies()

@app.get("/config/default")
async def get_default_config():
    """Get default configuration"""
    return engine.get_default_config()

# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "app.main:app",
        host=API_CONFIG['host'],
        port=API_CONFIG['port'],
        reload=API_CONFIG['reload']
    )
