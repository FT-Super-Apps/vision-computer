#!/usr/bin/env python3
"""
Celery Background Tasks
All heavy-duty processing tasks for concurrent execution
"""

from celery import Task
from app.celery_app import celery_app
import subprocess
import os
from datetime import datetime
from pathlib import Path
import pdfplumber
from fuzzywuzzy import fuzz
from docx import Document
import re
import sys

# Add parent directory to path
sys.path.append(str(Path(__file__).parent.parent))

from app.bypass_engine import BypassEngine

# Initialize bypass engine
engine = BypassEngine()


class ProgressTask(Task):
    """Base task with progress tracking"""

    def update_progress(self, current, total, message=""):
        """Update task progress"""
        self.update_state(
            state='PROGRESS',
            meta={
                'current': current,
                'total': total,
                'message': message,
                'percent': int((current / total) * 100) if total > 0 else 0
            }
        )


@celery_app.task(bind=True, base=ProgressTask, name='app.tasks.analyze_detect_flags_task')
def analyze_detect_flags_task(self, file_path: str, filename: str):
    """
    Background task: Detect flags dari Turnitin PDF

    Args:
        file_path: Path to uploaded PDF file
        filename: Original filename

    Returns:
        dict dengan flagged_items, total_pages, etc.
    """

    try:
        self.update_progress(1, 5, "Starting OCR process...")

        # Generate output path for OCR
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_pdf = f"temp/detect_ocr_{timestamp}_{filename}"

        self.update_progress(2, 5, "Running ocrmypdf --force-ocr...")

        # Run OCRmyPDF
        try:
            result = subprocess.run(
                ["ocrmypdf", file_path, output_pdf, "--force-ocr"],
                capture_output=True,
                text=True,
                timeout=300
            )
            if result.returncode != 0:
                raise Exception(f"OCR failed: {result.stderr}")
        except subprocess.TimeoutExpired:
            raise Exception("OCR timeout after 5 minutes")

        self.update_progress(3, 5, "Detecting highlights...")

        # Detect highlights dan extract flagged text
        flagged_items = []
        highlight_count = 0
        total_pages = 0

        with pdfplumber.open(file_path) as pdf_orig, pdfplumber.open(output_pdf) as pdf_ocr:
            total_pages = len(pdf_orig.pages)

            for page_num in range(total_pages):
                page_orig = pdf_orig.pages[page_num]
                page_ocr = pdf_ocr.pages[page_num]

                # Update progress per page
                if page_num % 5 == 0:
                    self.update_progress(
                        3 + (page_num / total_pages),
                        5,
                        f"Processing page {page_num + 1}/{total_pages}..."
                    )

                # Get colored rectangles (highlights)
                colored_rects = [
                    rect for rect in page_orig.rects
                    if rect['fill'] and rect['non_stroking_color'] and
                    rect['non_stroking_color'] not in [(0,0,0), (1,1,1)]
                ]

                if colored_rects:
                    highlight_count += len(colored_rects)

                    # Extract text dari highlighted areas
                    for rect in colored_rects:
                        bbox = (rect['x0'], rect['top'], rect['x1'], rect['bottom'])
                        cropped = page_ocr.within_bbox(bbox)
                        text = cropped.extract_text()

                        if text and len(text.strip()) > 3:
                            clean_text = text.strip()
                            if clean_text not in flagged_items:
                                flagged_items.append(clean_text)

        self.update_progress(5, 5, "Detection complete!")

        # Cleanup
        try:
            os.remove(file_path)
            os.remove(output_pdf)
        except:
            pass

        return {
            "success": True,
            "filename": filename,
            "total_pages": total_pages,
            "total_highlights": highlight_count,
            "flagged_items": flagged_items,
            "total_flags": len(flagged_items),
            "method": "Highlight Detection + OCRmyPDF"
        }

    except Exception as e:
        # Cleanup on error
        try:
            os.remove(file_path)
            if 'output_pdf' in locals():
                os.remove(output_pdf)
        except:
            pass

        raise Exception(f"Task failed: {str(e)}")


@celery_app.task(bind=True, base=ProgressTask, name='app.tasks.match_flags_task')
def match_flags_task(self, turnitin_pdf_path: str, original_doc_path: str,
                     turnitin_filename: str, original_filename: str):
    """
    Background task: Match flagged items dengan original document

    Args:
        turnitin_pdf_path: Path to Turnitin PDF
        original_doc_path: Path to original document
        turnitin_filename: Original Turnitin filename
        original_filename: Original document filename

    Returns:
        dict dengan matched_items, unmatched_items, etc.
    """

    try:
        self.update_progress(1, 4, "Extracting flags from Turnitin PDF...")

        # Step 1: Extract flagged items dari Turnitin PDF
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_pdf = f"temp/match_{timestamp}_turnitin_ocr.pdf"

        # Run OCRmyPDF
        subprocess.run(
            ["ocrmypdf", turnitin_pdf_path, output_pdf, "--force-ocr"],
            capture_output=True,
            text=True,
            timeout=300
        )

        # Extract flagged items
        flagged_items = []
        with pdfplumber.open(turnitin_pdf_path) as pdf_orig, pdfplumber.open(output_pdf) as pdf_ocr:
            for page_num in range(len(pdf_orig.pages)):
                page_orig = pdf_orig.pages[page_num]
                page_ocr = pdf_ocr.pages[page_num]

                colored_rects = [
                    rect for rect in page_orig.rects
                    if rect['fill'] and rect['non_stroking_color'] and
                    rect['non_stroking_color'] not in [(0,0,0), (1,1,1)]
                ]

                for rect in colored_rects:
                    bbox = (rect['x0'], rect['top'], rect['x1'], rect['bottom'])
                    cropped = page_ocr.within_bbox(bbox)
                    text = cropped.extract_text()

                    if text and len(text.strip()) > 3:
                        clean_text = text.strip()
                        if clean_text not in flagged_items:
                            flagged_items.append(clean_text)

        self.update_progress(2, 4, f"Extracting text from original document ({original_filename})...")

        # Step 2: Extract text dari original document
        original_text = ""
        doc_type = ""

        if original_filename.lower().endswith('.docx'):
            doc_type = "DOCX"
            doc = Document(original_doc_path)
            original_text = "\n".join([para.text for para in doc.paragraphs])

        elif original_filename.lower().endswith('.pdf'):
            doc_type = "PDF"
            temp_pdf_ocr = f"temp/match_{timestamp}_original_ocr.pdf"

            try:
                subprocess.run(
                    ["ocrmypdf", original_doc_path, temp_pdf_ocr, "--force-ocr"],
                    capture_output=True,
                    text=True,
                    timeout=300
                )
                pdf_to_read = temp_pdf_ocr
            except:
                pdf_to_read = original_doc_path

            with pdfplumber.open(pdf_to_read) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        original_text += text + "\n"

            try:
                os.remove(temp_pdf_ocr)
            except:
                pass

        elif original_filename.lower().endswith('.txt'):
            doc_type = "TXT"
            with open(original_doc_path, 'r', encoding='utf-8', errors='ignore') as f:
                original_text = f.read()

        self.update_progress(3, 4, f"Fuzzy matching {len(flagged_items)} items...")

        # Step 3: Fuzzy match
        matched_items = []
        unmatched_items = []

        for idx, flagged_text in enumerate(flagged_items):
            # Update progress every 10 items
            if idx % 10 == 0:
                self.update_progress(
                    3 + (idx / len(flagged_items) * 0.9),
                    4,
                    f"Matching {idx + 1}/{len(flagged_items)} items..."
                )

            best_match = None
            best_score = 0
            chunk_size = len(flagged_text) + 50

            for i in range(0, len(original_text) - len(flagged_text) + 1, 10):
                chunk = original_text[i:i + chunk_size]
                score = fuzz.token_set_ratio(flagged_text.lower(), chunk.lower())

                if score > best_score:
                    best_score = score
                    best_match = chunk

            if best_score >= 80:
                matched_items.append({
                    "flagged_text": flagged_text,
                    "matched_text": best_match[:100] + "..." if best_match and len(best_match) > 100 else best_match,
                    "similarity_score": best_score
                })
            else:
                unmatched_items.append({
                    "flagged_text": flagged_text,
                    "best_score": best_score
                })

        self.update_progress(4, 4, "Match complete!")

        # Cleanup
        try:
            os.remove(turnitin_pdf_path)
            os.remove(original_doc_path)
            os.remove(output_pdf)
        except:
            pass

        return {
            "success": True,
            "turnitin_filename": turnitin_filename,
            "original_filename": original_filename,
            "original_doc_type": doc_type,
            "total_flagged": len(flagged_items),
            "total_matched": len(matched_items),
            "total_unmatched": len(unmatched_items),
            "match_percentage": round((len(matched_items) / len(flagged_items) * 100) if flagged_items else 0, 2),
            "matched_items": matched_items,
            "unmatched_items": unmatched_items,
            "method": "Highlight Detection + Fuzzy Matching"
        }

    except Exception as e:
        # Cleanup on error
        try:
            os.remove(turnitin_pdf_path)
            os.remove(original_doc_path)
            if 'output_pdf' in locals():
                os.remove(output_pdf)
        except:
            pass

        raise Exception(f"Task failed: {str(e)}")


@celery_app.task(bind=True, base=ProgressTask, name='app.tasks.bypass_matched_flags_task')
def bypass_matched_flags_task(self, original_doc_path: str, flagged_text: str,
                              original_filename: str, homoglyph_density: float = 0.95,
                              invisible_density: float = 0.40):
    """
    Background task: Apply bypass ke matched flags

    Args:
        original_doc_path: Path to original DOCX
        flagged_text: Newline-separated list of flagged texts
        original_filename: Original filename
        homoglyph_density: Homoglyph density (0.0-1.0)
        invisible_density: Invisible character density (0.0-1.0)

    Returns:
        dict dengan processed_flags, output_file, etc.
    """

    try:
        self.update_progress(1, 4, "Loading document...")

        # Parse flagged texts
        flagged_list = [line.strip() for line in flagged_text.split('\n') if line.strip()]

        # Load document
        doc = Document(original_doc_path)

        self.update_progress(2, 4, f"Processing {len(flagged_list)} flagged items...")

        # Track modifications
        total_replacements = 0
        processed_flags = []

        # Process each flagged text
        for idx, flagged in enumerate(flagged_list):
            if not flagged:
                continue

            # Update progress
            if idx % 5 == 0:
                self.update_progress(
                    2 + (idx / len(flagged_list) * 1.5),
                    4,
                    f"Modifying {idx + 1}/{len(flagged_list)} items..."
                )

            replacements_for_this = 0

            # Search and replace in paragraphs
            for para in doc.paragraphs:
                if flagged.lower() in para.text.lower():
                    pattern = re.compile(re.escape(flagged), re.IGNORECASE)

                    if pattern.search(para.text):
                        modified_version = engine.apply_combined_bypass(
                            flagged,
                            homoglyph_density=homoglyph_density,
                            invisible_density=invisible_density
                        )

                        for run in para.runs:
                            if flagged.lower() in run.text.lower():
                                matches = list(pattern.finditer(run.text))
                                if matches:
                                    replacements_for_this += len(matches)
                                    total_replacements += len(matches)
                                    run.text = pattern.sub(modified_version, run.text)

            # Also process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if flagged.lower() in para.text.lower():
                                pattern = re.compile(re.escape(flagged), re.IGNORECASE)
                                modified_version = engine.apply_combined_bypass(
                                    flagged,
                                    homoglyph_density=homoglyph_density,
                                    invisible_density=invisible_density
                                )

                                for run in para.runs:
                                    if flagged.lower() in run.text.lower():
                                        if pattern.search(run.text):
                                            total_replacements += 1
                                            run.text = pattern.sub(modified_version, run.text)

            if replacements_for_this > 0:
                processed_flags.append({
                    "flagged_text": flagged,
                    "replacements_made": replacements_for_this,
                    "modified_version": engine.apply_combined_bypass(flagged, homoglyph_density, invisible_density)[:50] + "..."
                })

        self.update_progress(3.5, 4, "Saving modified document...")

        # Save modified document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"outputs/modified_bypass_{timestamp}.docx"
        doc.save(output_filename)

        self.update_progress(4, 4, "Bypass complete!")

        # Cleanup
        try:
            os.remove(original_doc_path)
        except:
            pass

        return {
            "success": True,
            "original_filename": original_filename,
            "total_flagged_items": len(flagged_list),
            "total_replacements": total_replacements,
            "processed_items": len(processed_flags),
            "processed_flags": processed_flags,
            "output_file": output_filename,
            "homoglyph_density": homoglyph_density,
            "invisible_density": invisible_density,
            "method": "Combined Homoglyph + Invisible Characters"
        }

    except Exception as e:
        # Cleanup on error
        try:
            os.remove(original_doc_path)
        except:
            pass

        raise Exception(f"Task failed: {str(e)}")


@celery_app.task(bind=True, base=ProgressTask, name='app.tasks.process_document_unified_task')
def process_document_unified_task(self, turnitin_pdf_path: str, original_doc_path: str,
                                  turnitin_filename: str, original_filename: str,
                                  homoglyph_density: float = 0.95, invisible_density: float = 0.40):
    """
    Unified background task: Analyze → Match → Bypass in one process

    Args:
        turnitin_pdf_path: Path to Turnitin PDF
        original_doc_path: Path to original document (DOCX)
        turnitin_filename: Turnitin PDF filename
        original_filename: Original document filename
        homoglyph_density: Homoglyph density (0.0-1.0), default 0.95
        invisible_density: Invisible character density (0.0-1.0), default 0.40

    Returns:
        dict dengan complete results: flagged items, matched items, output file, stats
    """

    TOTAL_STEPS = 13  # 5 (analyze) + 4 (match) + 4 (bypass)

    try:
        # ============================================
        # PHASE 1: ANALYZE/DETECT FLAGS (Steps 1-5)
        # ============================================
        self.update_progress(1, TOTAL_STEPS, "Phase 1/3: Starting OCR process...")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_pdf = f"temp/unified_{timestamp}_turnitin_ocr.pdf"

        self.update_progress(2, TOTAL_STEPS, "Phase 1/3: Running OCRmyPDF on Turnitin PDF...")

        # Run OCRmyPDF
        try:
            result = subprocess.run(
                ["ocrmypdf", turnitin_pdf_path, output_pdf, "--force-ocr"],
                capture_output=True,
                text=True,
                timeout=300
            )
            if result.returncode != 0:
                raise Exception(f"OCR failed: {result.stderr}")
        except subprocess.TimeoutExpired:
            raise Exception("OCR timeout after 5 minutes")

        self.update_progress(3, TOTAL_STEPS, "Phase 1/3: Detecting highlights...")

        # Detect highlights and extract flagged text
        flagged_items = []
        highlight_count = 0
        total_pages = 0

        with pdfplumber.open(turnitin_pdf_path) as pdf_orig, pdfplumber.open(output_pdf) as pdf_ocr:
            total_pages = len(pdf_orig.pages)

            for page_num in range(total_pages):
                page_orig = pdf_orig.pages[page_num]
                page_ocr = pdf_ocr.pages[page_num]

                # Update progress per page
                if page_num % 5 == 0:
                    progress = 3 + (page_num / total_pages * 2)  # Steps 3-5
                    self.update_progress(
                        progress,
                        TOTAL_STEPS,
                        f"Phase 1/3: Processing page {page_num + 1}/{total_pages}..."
                    )

                # Get colored rectangles (highlights)
                colored_rects = [
                    rect for rect in page_orig.rects
                    if rect['fill'] and rect['non_stroking_color'] and
                    rect['non_stroking_color'] not in [(0,0,0), (1,1,1)]
                ]

                if colored_rects:
                    highlight_count += len(colored_rects)

                    # Extract text from highlighted areas
                    for rect in colored_rects:
                        bbox = (rect['x0'], rect['top'], rect['x1'], rect['bottom'])
                        cropped = page_ocr.within_bbox(bbox)
                        text = cropped.extract_text()

                        if text and len(text.strip()) > 3:
                            clean_text = text.strip()
                            if clean_text not in flagged_items:
                                flagged_items.append(clean_text)

        self.update_progress(5, TOTAL_STEPS, f"Phase 1/3: Detected {len(flagged_items)} unique flags!")

        # ============================================
        # PHASE 2: MATCH FLAGS (Steps 6-9)
        # ============================================
        self.update_progress(6, TOTAL_STEPS, f"Phase 2/3: Extracting text from {original_filename}...")

        # Extract text from original document
        original_text = ""
        doc_type = ""

        if original_filename.lower().endswith('.docx'):
            doc_type = "DOCX"
            doc = Document(original_doc_path)
            original_text = "\n".join([para.text for para in doc.paragraphs])

        elif original_filename.lower().endswith('.pdf'):
            doc_type = "PDF"
            temp_pdf_ocr = f"temp/unified_{timestamp}_original_ocr.pdf"

            try:
                subprocess.run(
                    ["ocrmypdf", original_doc_path, temp_pdf_ocr, "--force-ocr"],
                    capture_output=True,
                    text=True,
                    timeout=300
                )
                pdf_to_read = temp_pdf_ocr
            except:
                pdf_to_read = original_doc_path

            with pdfplumber.open(pdf_to_read) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        original_text += text + "\n"

            try:
                os.remove(temp_pdf_ocr)
            except:
                pass

        elif original_filename.lower().endswith('.txt'):
            doc_type = "TXT"
            with open(original_doc_path, 'r', encoding='utf-8', errors='ignore') as f:
                original_text = f.read()

        self.update_progress(7, TOTAL_STEPS, f"Phase 2/3: Fuzzy matching {len(flagged_items)} items...")

        # Fuzzy match
        matched_items = []
        unmatched_items = []

        for idx, flagged_text in enumerate(flagged_items):
            # Update progress every 10 items
            if idx % 10 == 0:
                progress = 7 + (idx / len(flagged_items) * 2)  # Steps 7-9
                self.update_progress(
                    progress,
                    TOTAL_STEPS,
                    f"Phase 2/3: Matching {idx + 1}/{len(flagged_items)}..."
                )

            best_match = None
            best_score = 0
            chunk_size = len(flagged_text) + 50

            for i in range(0, len(original_text) - len(flagged_text) + 1, 10):
                chunk = original_text[i:i + chunk_size]
                score = fuzz.token_set_ratio(flagged_text.lower(), chunk.lower())

                if score > best_score:
                    best_score = score
                    best_match = chunk

            if best_score >= 80:
                matched_items.append({
                    "flagged_text": flagged_text,
                    "matched_text": best_match[:100] + "..." if best_match and len(best_match) > 100 else best_match,
                    "similarity_score": best_score
                })
            else:
                unmatched_items.append({
                    "flagged_text": flagged_text,
                    "best_score": best_score
                })

        self.update_progress(9, TOTAL_STEPS, f"Phase 2/3: Matched {len(matched_items)}/{len(flagged_items)} items!")

        # ============================================
        # PHASE 3: BYPASS MATCHED FLAGS (Steps 10-13)
        # ============================================
        self.update_progress(10, TOTAL_STEPS, "Phase 3/3: Loading document for bypass...")

        # Only process DOCX for bypass
        if not original_filename.lower().endswith('.docx'):
            raise Exception("Bypass only supports DOCX files. Please provide original document in DOCX format.")

        # Reload document for bypass
        doc = Document(original_doc_path)

        self.update_progress(11, TOTAL_STEPS, f"Phase 3/3: Applying bypass to {len(matched_items)} matched items...")

        # Track modifications
        total_replacements = 0
        processed_flags = []

        # Only bypass matched items
        matched_texts = [item['flagged_text'] for item in matched_items]

        for idx, flagged in enumerate(matched_texts):
            if not flagged:
                continue

            # Update progress
            if idx % 5 == 0:
                progress = 11 + (idx / len(matched_texts) * 1.5)  # Steps 11-12.5
                self.update_progress(
                    progress,
                    TOTAL_STEPS,
                    f"Phase 3/3: Modifying {idx + 1}/{len(matched_texts)}..."
                )

            replacements_for_this = 0

            # Search and replace in paragraphs
            for para in doc.paragraphs:
                if flagged.lower() in para.text.lower():
                    pattern = re.compile(re.escape(flagged), re.IGNORECASE)

                    if pattern.search(para.text):
                        modified_version = engine.apply_combined_bypass(
                            flagged,
                            homoglyph_density=homoglyph_density,
                            invisible_density=invisible_density
                        )

                        for run in para.runs:
                            if flagged.lower() in run.text.lower():
                                matches = list(pattern.finditer(run.text))
                                if matches:
                                    replacements_for_this += len(matches)
                                    total_replacements += len(matches)
                                    run.text = pattern.sub(modified_version, run.text)

            # Also process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if flagged.lower() in para.text.lower():
                                pattern = re.compile(re.escape(flagged), re.IGNORECASE)
                                modified_version = engine.apply_combined_bypass(
                                    flagged,
                                    homoglyph_density=homoglyph_density,
                                    invisible_density=invisible_density
                                )

                                for run in para.runs:
                                    if flagged.lower() in run.text.lower():
                                        if pattern.search(run.text):
                                            replacements_for_this += 1
                                            total_replacements += 1
                                            run.text = pattern.sub(modified_version, run.text)

            if replacements_for_this > 0:
                processed_flags.append({
                    "flagged_text": flagged,
                    "replacements_made": replacements_for_this,
                    "modified_version": engine.apply_combined_bypass(flagged, homoglyph_density, invisible_density)[:50] + "..."
                })

        self.update_progress(12.5, TOTAL_STEPS, "Phase 3/3: Saving modified document...")

        # Save modified document
        output_filename = f"outputs/unified_bypass_{timestamp}_{original_filename}"
        doc.save(output_filename)

        self.update_progress(13, TOTAL_STEPS, "Complete! Document processed successfully.")

        # Cleanup
        try:
            os.remove(turnitin_pdf_path)
            os.remove(original_doc_path)
            os.remove(output_pdf)
        except:
            pass

        # Return comprehensive results
        return {
            "success": True,
            "turnitin_filename": turnitin_filename,
            "original_filename": original_filename,
            "original_doc_type": doc_type,

            # Phase 1 results
            "total_pages": total_pages,
            "total_highlights": highlight_count,
            "total_flags": len(flagged_items),
            "flagged_items": flagged_items,

            # Phase 2 results
            "total_matched": len(matched_items),
            "total_unmatched": len(unmatched_items),
            "match_percentage": round((len(matched_items) / len(flagged_items) * 100) if flagged_items else 0, 2),
            "matched_items": matched_items,
            "unmatched_items": unmatched_items,

            # Phase 3 results
            "total_replacements": total_replacements,
            "processed_items": len(processed_flags),
            "processed_flags": processed_flags,
            "output_file": output_filename,

            # Settings
            "homoglyph_density": homoglyph_density,
            "invisible_density": invisible_density,

            "method": "Unified: Detect → Match → Bypass"
        }

    except Exception as e:
        # Cleanup on error
        try:
            os.remove(turnitin_pdf_path)
            os.remove(original_doc_path)
            if 'output_pdf' in locals():
                os.remove(output_pdf)
        except:
            pass

        raise Exception(f"Unified task failed: {str(e)}")
