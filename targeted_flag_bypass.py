#!/usr/bin/env python3
"""
Targeted Turnitin Flag Bypass
Hanya memodifikasi kata/frasa yang ter-flag oleh Turnitin
Menggunakan kombinasi: Homoglyphs + Invisible Characters
"""

import random
from docx import Document
from docx.shared import Pt

# ============================================================================
# CONFIGURATION
# ============================================================================

# Cyrillic homoglyphs (visual identical)
HOMOGLYPHS = {
    'A': 'А', 'B': 'В', 'C': 'С', 'E': 'Е', 'H': 'Н',
    'I': 'І', 'K': 'К', 'M': 'М', 'O': 'О', 'P': 'Р',
    'T': 'Т', 'X': 'Х', 'Y': 'У',
    'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с',
    'x': 'х', 'y': 'у', 'i': 'і', 'k': 'к', 'u': 'υ',
    'd': 'ԁ', 'g': 'ց', 'j': 'ј', 'l': 'l', 'n': 'п',
    's': 'ѕ', 'v': 'ѵ', 'w': 'ԝ', 'z': 'ᴢ'
}

# Invisible characters (zero-width)
INVISIBLE_CHARS = [
    '\u200B',  # Zero-width space
    '\u200C',  # Zero-width non-joiner
    '\u200D',  # Zero-width joiner
    '\uFEFF',  # Zero-width no-break space
]

# ============================================================================
# LOAD FLAGGED PHRASES
# ============================================================================

def load_flagged_phrases(flag_file='flag.txt'):
    """Load kata/frasa yang ter-flag dari file"""
    with open(flag_file, 'r', encoding='utf-8') as f:
        phrases = [line.strip() for line in f if line.strip()]

    # Sort by length (longest first) untuk avoid partial replacement
    phrases.sort(key=len, reverse=True)

    print(f"✅ Loaded {len(phrases)} flagged phrases")
    return phrases

# ============================================================================
# BYPASS TECHNIQUES
# ============================================================================

def apply_homoglyphs(text, density=0.80):
    """Apply homoglyphs replacement - SMART selection for natural look"""
    if not text:
        return text

    result = list(text)
    replaceable = [i for i, c in enumerate(text) if c in HOMOGLYPHS]

    if not replaceable:
        return text

    # SMART SELECTION: Prioritize less obvious characters
    # Characters that look MOST similar (hard to detect)
    high_priority = {'a', 'e', 'o', 'c', 'p', 'x', 'A', 'E', 'O', 'C', 'P', 'X'}

    # Split into priority groups
    high_priority_pos = [i for i in replaceable if result[i] in high_priority]
    other_pos = [i for i in replaceable if result[i] not in high_priority]

    # Replace more from high priority (looks more natural)
    num_high = int(len(high_priority_pos) * min(density * 1.2, 1.0))
    num_other = int(len(other_pos) * density * 0.5)

    selected_pos = []
    if high_priority_pos:
        selected_pos.extend(random.sample(high_priority_pos, min(num_high, len(high_priority_pos))))
    if other_pos and num_other > 0:
        selected_pos.extend(random.sample(other_pos, min(num_other, len(other_pos))))

    # Apply replacements
    for pos in selected_pos:
        result[pos] = HOMOGLYPHS[result[pos]]

    return ''.join(result)

def apply_invisible_chars(text, density=0.30):
    """Insert invisible characters - SMART placement for natural look"""
    if not text or len(text) < 3:
        return text

    result = list(text)

    # SMART PLACEMENT: Insert at word boundaries (more natural)
    # Find spaces (word boundaries)
    word_boundaries = [i for i, c in enumerate(text) if c == ' ']

    if not word_boundaries:
        return text

    # Insert invisible chars near word boundaries (less detectable)
    num_insert = max(1, int(len(word_boundaries) * density))
    selected_boundaries = random.sample(word_boundaries, min(num_insert, len(word_boundaries)))

    # Sort reverse to maintain positions
    for pos in sorted(selected_boundaries, reverse=True):
        # Insert right after space (between words)
        if pos + 1 < len(result):
            invisible_char = random.choice(INVISIBLE_CHARS)
            result.insert(pos + 1, invisible_char)

    return ''.join(result)

def apply_combined_bypass(text, homoglyph_density=0.50, invisible_density=0.15):
    """Kombinasi homoglyphs + invisible characters - NATURAL settings"""
    # Step 1: Apply homoglyphs (smart selection prioritizes natural-looking chars)
    text = apply_homoglyphs(text, density=homoglyph_density)

    # Step 2: Apply invisible characters (placed at word boundaries)
    text = apply_invisible_chars(text, density=invisible_density)

    return text

# ============================================================================
# TARGETED REPLACEMENT
# ============================================================================

def targeted_replace_in_text(text, flagged_phrases, homoglyph_density=0.80, invisible_density=0.30):
    """Replace hanya kata/frasa yang ter-flag"""
    if not text:
        return text, 0

    modified_text = text
    replacement_count = 0

    for phrase in flagged_phrases:
        if phrase in modified_text:
            # Apply bypass technique
            bypassed_phrase = apply_combined_bypass(phrase, homoglyph_density, invisible_density)

            # Replace all occurrences
            modified_text = modified_text.replace(phrase, bypassed_phrase)
            replacement_count += 1

    return modified_text, replacement_count

# ============================================================================
# MAIN PROCESSING
# ============================================================================

def process_document(input_docx='original.docx',
                     output_docx='original_targeted_bypass.docx',
                     flag_file='flag.txt',
                     homoglyph_density=0.50,
                     invisible_density=0.15):
    """Process dokumen dengan targeted bypass"""

    print(f"\n{'='*70}")
    print(f"🎯 TARGETED TURNITIN FLAG BYPASS")
    print(f"{'='*70}")
    print(f"Input       : {input_docx}")
    print(f"Output      : {output_docx}")
    print(f"Flag file   : {flag_file}")
    print(f"Homoglyph   : {int(homoglyph_density*100)}%")
    print(f"Invisible   : {int(invisible_density*100)}%")
    print(f"{'='*70}\n")

    # Check if output file exists
    import os
    if os.path.exists(output_docx):
        print(f"⚠️  WARNING: {output_docx} already exists!")

        # Generate new filename with timestamp
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = output_docx.rsplit('.', 1)[0]
        output_docx = f"{base_name}_{timestamp}.docx"
        print(f"✅ Creating new file: {output_docx}\n")

    # Load flagged phrases
    flagged_phrases = load_flagged_phrases(flag_file)

    print(f"\n📋 Flagged phrases preview:")
    for i, phrase in enumerate(flagged_phrases[:5], 1):
        print(f"   {i}. {phrase[:60]}{'...' if len(phrase) > 60 else ''}")
    if len(flagged_phrases) > 5:
        print(f"   ... and {len(flagged_phrases) - 5} more")
    print()

    # Load document (read from input, never modify input)
    doc = Document(input_docx)

    total_replacements = 0
    modified_paragraphs = 0

    # Process each paragraph
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue

        para_modified = False

        # Process each run to preserve formatting
        for run in para.runs:
            if run.text:
                modified_text, count = targeted_replace_in_text(
                    run.text,
                    flagged_phrases,
                    homoglyph_density,
                    invisible_density
                )

                if count > 0:
                    run.text = modified_text
                    total_replacements += count
                    para_modified = True

        if para_modified:
            modified_paragraphs += 1

            if modified_paragraphs <= 3:
                print(f"✅ Para {i+1}: {total_replacements} flagged phrases replaced")

    # Save document
    doc.save(output_docx)

    print(f"\n{'='*70}")
    print(f"✅ COMPLETE!")
    print(f"{'='*70}")
    print(f"Modified paragraphs    : {modified_paragraphs}")
    print(f"Total replacements     : {total_replacements}")
    print(f"Output saved to        : {output_docx}")
    print(f"\n🎯 Strategi:")
    print(f"   - Homoglyphs: {int(homoglyph_density*100)}% karakter diganti (Cyrillic)")
    print(f"   - Invisible : {int(invisible_density*100)}% invisible chars inserted")
    print(f"   - Font/formatting tetap dipertahankan")
    print(f"{'='*70}\n")

# ============================================================================
# ANALYZE MODE
# ============================================================================

def analyze_flags(input_docx='original.docx', flag_file='flag.txt'):
    """Analisis berapa banyak flagged phrases di dokumen"""

    print(f"\n{'='*70}")
    print(f"🔍 ANALYZE FLAGGED PHRASES IN DOCUMENT")
    print(f"{'='*70}\n")

    flagged_phrases = load_flagged_phrases(flag_file)
    doc = Document(input_docx)

    # Get all text
    full_text = '\n'.join([para.text for para in doc.paragraphs])

    print(f"📊 Analysis Results:\n")

    found_count = 0
    total_occurrences = 0

    for phrase in flagged_phrases:
        count = full_text.count(phrase)
        if count > 0:
            found_count += 1
            total_occurrences += count
            print(f"   ✓ '{phrase[:50]}...' → {count} occurrence(s)")

    print(f"\n{'='*70}")
    print(f"Summary:")
    print(f"   - Total flagged phrases : {len(flagged_phrases)}")
    print(f"   - Found in document     : {found_count}")
    print(f"   - Total occurrences     : {total_occurrences}")
    print(f"{'='*70}\n")

# ============================================================================
# DETECTION MODE - Verify Bypass Success
# ============================================================================

def detect_bypass_success(original_docx='original.docx',
                          bypassed_docx='original_natural.docx',
                          flag_file='flag.txt'):
    """Compare original vs bypassed untuk verify bypass berhasil"""

    print(f"\n{'='*70}")
    print(f"🔬 BYPASS SUCCESS DETECTION")
    print(f"{'='*70}")
    print(f"Original  : {original_docx}")
    print(f"Bypassed  : {bypassed_docx}")
    print(f"Flag file : {flag_file}")
    print(f"{'='*70}\n")

    flagged_phrases = load_flagged_phrases(flag_file)

    # Load both documents
    doc_original = Document(original_docx)
    doc_bypassed = Document(bypassed_docx)

    text_original = '\n'.join([para.text for para in doc_original.paragraphs])
    text_bypassed = '\n'.join([para.text for para in doc_bypassed.paragraphs])

    print(f"📊 Detection Results:\n")
    print(f"{'Phrase':<50} {'Original':<12} {'Bypassed':<12} {'Status':<10}")
    print(f"{'-'*50} {'-'*12} {'-'*12} {'-'*10}")

    bypassed_count = 0
    still_detected = 0

    for phrase in flagged_phrases:
        count_orig = text_original.count(phrase)
        count_bypass = text_bypassed.count(phrase)

        if count_orig > 0:  # Only show phrases that existed in original
            status = "✅ HIDDEN" if count_bypass == 0 else "❌ DETECTED"

            if count_bypass == 0:
                bypassed_count += 1
            else:
                still_detected += 1

            phrase_short = phrase[:47] + '...' if len(phrase) > 50 else phrase
            print(f"{phrase_short:<50} {count_orig:<12} {count_bypass:<12} {status:<10}")

    print(f"\n{'='*70}")
    print(f"📈 BYPASS SUCCESS RATE:")
    print(f"{'='*70}")

    total_original = sum(text_original.count(p) for p in flagged_phrases)
    total_bypassed = sum(text_bypassed.count(p) for p in flagged_phrases)

    if total_original > 0:
        success_rate = (bypassed_count / (bypassed_count + still_detected)) * 100
        reduction_rate = ((total_original - total_bypassed) / total_original) * 100
    else:
        success_rate = 0
        reduction_rate = 0

    print(f"   Phrases successfully hidden  : {bypassed_count}")
    print(f"   Phrases still detected       : {still_detected}")
    print(f"   Success rate                 : {success_rate:.1f}%")
    print(f"   Occurrence reduction         : {reduction_rate:.1f}%")
    print(f"\n   Original occurrences         : {total_original}")
    print(f"   Bypassed occurrences         : {total_bypassed}")

    if success_rate == 100:
        print(f"\n🎉 PERFECT! All flagged phrases are hidden!")
    elif success_rate >= 80:
        print(f"\n✅ EXCELLENT! Most phrases are hidden.")
    elif success_rate >= 50:
        print(f"\n⚠️  MODERATE. Consider increasing density.")
    else:
        print(f"\n❌ LOW SUCCESS. Need more aggressive strategy.")

    print(f"{'='*70}\n")

# ============================================================================
# MAIN
# ============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == 'analyze':
        # Analyze mode - cek flagged phrases di dokumen
        input_file = sys.argv[2] if len(sys.argv) > 2 else 'original.docx'
        flag_file = sys.argv[3] if len(sys.argv) > 3 else 'flag.txt'
        analyze_flags(input_file, flag_file)

    elif len(sys.argv) > 1 and sys.argv[1] == 'detect':
        # Detection mode - compare original vs bypassed
        original_file = sys.argv[2] if len(sys.argv) > 2 else 'original.docx'
        bypassed_file = sys.argv[3] if len(sys.argv) > 3 else 'original_natural.docx'
        flag_file = sys.argv[4] if len(sys.argv) > 4 else 'flag.txt'
        detect_bypass_success(original_file, bypassed_file, flag_file)

    else:
        # Process mode - jalankan bypass
        input_file = sys.argv[1] if len(sys.argv) > 1 else 'original.docx'
        output_file = sys.argv[2] if len(sys.argv) > 2 else 'original_targeted_bypass.docx'
        flag_file = sys.argv[3] if len(sys.argv) > 3 else 'flag.txt'

        # Optional: custom densities (default: natural settings)
        homoglyph_d = float(sys.argv[4]) if len(sys.argv) > 4 else 0.50
        invisible_d = float(sys.argv[5]) if len(sys.argv) > 5 else 0.15

        process_document(input_file, output_file, flag_file, homoglyph_d, invisible_d)
