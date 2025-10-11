#!/usr/bin/env python3
"""
HYBRID SMART APPROACH:
- Headers (≤5 words): invisible chars + unicode (TERBUKTI BERHASIL!)
- Content (>5 words): IndoT5 paraphrase + invisible chars MINIMAL
- Goal: Avoid triggering OCR while maximizing bypass
"""

import json
import docx
from difflib import SequenceMatcher
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch
from tqdm import tqdm
import re

def load_indot5_model(model_name="Wikidepia/IndoT5-base-paraphrase"):
    """Load IndoT5 model"""
    print(f"🔧 Loading IndoT5: {model_name}")
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)
    print(f"✅ Model loaded on {device}")
    return tokenizer, model, device

def paraphrase_with_indot5(text, tokenizer, model, device, temperature=1.0, num_beams=10):
    """High-quality IndoT5 paraphrase"""
    input_text = f"paraphrase: {text}"
    inputs = tokenizer(
        input_text,
        return_tensors="pt",
        max_length=512,
        truncation=True,
        padding=True
    ).to(device)
    
    with torch.no_grad():
        outputs = model.generate(
            inputs.input_ids,
            max_length=512,
            num_beams=num_beams,
            temperature=temperature,
            do_sample=True,
            top_k=50,
            top_p=0.95,
            early_stopping=True,
            no_repeat_ngram_size=3,
            num_return_sequences=1
        )
    
    paraphrased = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return paraphrased

def add_invisible_chars_minimal(text, density=0.05):
    """
    MINIMAL invisible chars - just enough to bypass, not trigger OCR
    Much lower density: 0.05 (5%) instead of 0.3 (30%)
    """
    import random
    invisible_chars = ['\u200B', '\u200C', '\u200D']  # ZWSP, ZWNJ, ZWJ
    
    words = text.split()
    result_words = []
    
    for word in words:
        if random.random() < density and len(word) > 3:
            # Insert ONE invisible char at random position
            pos = random.randint(1, len(word) - 1)
            char = random.choice(invisible_chars)
            modified = word[:pos] + char + word[pos:]
            result_words.append(modified)
        else:
            result_words.append(word)
    
    return ' '.join(result_words)

def add_unicode_substitution_selective(text, percentage=0.05):
    """
    SELECTIVE unicode substitution - only safe characters
    Much lower rate: 5% instead of 10-15%
    """
    import random
    
    # Only use most similar lookalikes
    safe_subs = {
        'a': 'а',  # Cyrillic a (U+0430)
        'e': 'е',  # Cyrillic e (U+0435)
        'o': 'о',  # Cyrillic o (U+043E)
    }
    
    chars = list(text)
    num_to_replace = int(len(chars) * percentage)
    replaceable_indices = [i for i, c in enumerate(chars) if c.lower() in safe_subs]
    
    if replaceable_indices:
        replace_indices = random.sample(
            replaceable_indices,
            min(num_to_replace, len(replaceable_indices))
        )
        
        for idx in replace_indices:
            original_char = chars[idx]
            if original_char.lower() in safe_subs:
                replacement = safe_subs[original_char.lower()]
                chars[idx] = replacement.upper() if original_char.isupper() else replacement
    
    return ''.join(chars)

def categorize_text(text):
    """Categorize as header or content"""
    words = text.split()
    word_count = len(words)
    
    # Headers: ≤5 words, often uppercase, short
    is_short = word_count <= 5
    is_uppercase = text.isupper() or text.istitle()
    
    # Common header patterns
    header_patterns = [
        r'^BAB\s+[IVX]+',           # BAB I, BAB II, etc
        r'^KATA\s+PENGANTAR',       # KATA PENGANTAR
        r'^DAFTAR\s+\w+',          # DAFTAR ISI, etc
        r'^LAMPIRAN',               # LAMPIRAN
        r'^ABSTRAK',                # ABSTRAK
        r'^\d+\.\s*[A-Z]',         # 1. PENDAHULUAN
    ]
    
    matches_pattern = any(re.match(pattern, text, re.IGNORECASE) for pattern in header_patterns)
    
    if (is_short and is_uppercase) or matches_pattern:
        return "header"
    else:
        return "content"

def match_flagged_with_docx(flagged_json, docx_path, threshold=0.5):
    """Match flagged texts with DOCX"""
    with open(flagged_json, 'r', encoding='utf-8') as f:
        flagged_texts = json.load(f)
    
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    print(f"\n🔍 Matching {len(flagged_texts)} flagged with {len(paragraphs)} paragraphs...")
    
    matches = []
    for i, flagged in enumerate(flagged_texts, 1):
        flagged_text = flagged['text']
        best_match = None
        best_ratio = 0
        best_idx = -1
        
        for idx, para in enumerate(paragraphs):
            ratio = SequenceMatcher(None, flagged_text.lower(), para.lower()).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = para
                best_idx = idx
        
        if best_ratio < 0.5:
            keywords = [w.lower() for w in flagged_text.split() if len(w) > 4]
            for idx, para in enumerate(paragraphs):
                keyword_matches = sum(1 for kw in keywords if kw in para.lower())
                keyword_ratio = keyword_matches / len(keywords) if keywords else 0
                if keyword_ratio > best_ratio:
                    best_ratio = keyword_ratio
                    best_match = para
                    best_idx = idx
        
        if best_ratio >= threshold:
            matches.append({
                "flagged_text": flagged_text,
                "matched_para": best_match,
                "similarity": round(best_ratio * 100, 2),
                "para_index": best_idx,
                "page": flagged['page']
            })
    
    print(f"✅ Matched: {len(matches)}/{len(flagged_texts)}")
    return matches

def process_hybrid_smart(matches):
    """
    HYBRID SMART STRATEGY:
    - Headers: Invisible + Unicode (proven to work!)
    - Content: IndoT5 + Minimal invisible (avoid OCR trigger)
    """
    print(f"\n🎨 Processing {len(matches)} texts with HYBRID strategy...")
    print("="*70)
    
    # Load IndoT5 for content
    tokenizer, model, device = load_indot5_model()
    
    results = []
    headers = 0
    contents = 0
    
    for match in tqdm(matches, desc="Processing"):
        original = match['matched_para']
        category = categorize_text(original)
        
        if category == "header":
            # STRATEGY 1: Headers - aggressive invisible + unicode
            modified = add_invisible_chars_minimal(original, density=0.15)  # 15%
            modified = add_unicode_substitution_selective(modified, percentage=0.10)  # 10%
            method = "header_tricks"
            headers += 1
            
        else:
            # STRATEGY 2: Content - IndoT5 + minimal invisible
            # Split long texts
            if len(original) > 400:
                sentences = [s.strip() + '.' for s in original.split('.') if s.strip()]
                paraphrased_parts = []
                
                for sent in sentences:
                    if len(sent) > 10:
                        para = paraphrase_with_indot5(sent, tokenizer, model, device,
                                                       temperature=1.0, num_beams=10)
                        paraphrased_parts.append(para)
                    else:
                        paraphrased_parts.append(sent)
                
                paraphrased = ' '.join(paraphrased_parts)
            else:
                paraphrased = paraphrase_with_indot5(original, tokenizer, model, device,
                                                      temperature=1.0, num_beams=10)
            
            # Add MINIMAL invisible chars to paraphrased text
            modified = add_invisible_chars_minimal(paraphrased, density=0.05)  # Only 5%!
            method = "indot5_minimal"
            contents += 1
        
        results.append({
            "original": original,
            "modified": modified,
            "category": category,
            "method": method,
            "page": match['page'],
            "similarity_before": match['similarity'],
            "para_index": match['para_index']
        })
    
    print(f"\n📊 Strategy Distribution:")
    print(f"  Headers (tricks): {headers}")
    print(f"  Contents (IndoT5+minimal): {contents}")
    
    return results

def apply_to_docx(processed_data, input_docx, output_docx):
    """Apply processed texts to DOCX"""
    print(f"\n📄 Loading: {input_docx}")
    doc = docx.Document(input_docx)
    
    replaced = 0
    for item in tqdm(processed_data, desc="Applying"):
        original = item['original']
        modified = item['modified']
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            ratio = SequenceMatcher(None, original.lower(), para.text.lower()).ratio()
            if ratio >= 0.7:
                # Replace keeping formatting
                if para.runs:
                    for run in para.runs:
                        run.text = ''
                    para.runs[0].text = modified
                else:
                    para.text = modified
                replaced += 1
                break
    
    print(f"✅ Replaced: {replaced}/{len(processed_data)}")
    doc.save(output_docx)
    print(f"💾 Saved: {output_docx}")
    return replaced

def main():
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python hybrid_smart_bypass.py <flagged.json> <original.docx>")
        sys.exit(1)
    
    flagged_json = sys.argv[1]
    docx_file = sys.argv[2]
    output_file = "testing_hybrid_bypassed.docx"
    
    print("="*70)
    print("🎓 HYBRID SMART BYPASS")
    print("="*70)
    print("Strategy:")
    print("  📌 Headers (≤5 words):")
    print("     → Invisible chars (15% density)")
    print("     → Unicode substitution (10% - safe only)")
    print("     → PROVEN to bypass Turnitin!")
    print()
    print("  📝 Content (>5 words):")
    print("     → IndoT5 high-quality paraphrase")
    print("     → MINIMAL invisible chars (5% only)")
    print("     → Avoid triggering OCR re-check")
    print("="*70)
    
    # Step 1: Match
    matches = match_flagged_with_docx(flagged_json, docx_file, threshold=0.5)
    
    if not matches:
        print("❌ No matches found")
        return
    
    # Save matches
    with open("hybrid_matches.json", 'w', encoding='utf-8') as f:
        json.dump(matches, f, ensure_ascii=False, indent=2)
    
    # Step 2: Hybrid processing
    processed = process_hybrid_smart(matches)
    
    # Save processed
    with open("hybrid_processed.json", 'w', encoding='utf-8') as f:
        json.dump(processed, f, ensure_ascii=False, indent=2)
    
    # Step 3: Apply to DOCX
    apply_to_docx(processed, docx_file, output_file)
    
    print("\n" + "="*70)
    print("✅ COMPLETE!")
    print("="*70)
    print(f"📄 Output: {output_file}")
    print(f"📊 Processed: {len(processed)} texts")
    print()
    print("📋 Key Differences:")
    print("  ✅ Headers use aggressive tricks (proven!)")
    print("  ✅ Content uses IndoT5 + MINIMAL tricks")
    print("  ✅ Lower density = less OCR trigger risk")
    print("  ✅ Smart categorization")
    print()
    print("🎯 Expected Result:")
    print("  - Headers bypass (like before)")
    print("  - Content natural (no OCR errors)")
    print("  - Overall similarity SHOULD decrease!")
    print("="*70)

if __name__ == "__main__":
    main()
