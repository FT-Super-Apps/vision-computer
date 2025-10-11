#!/usr/bin/env python3
"""
CLEAN APPROACH: Pure IndoT5 Paraphrase Only
No tricks, just high-quality natural paraphrasing
"""

import json
import docx
from difflib import SequenceMatcher
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch
from tqdm import tqdm

def load_indot5_model(model_name="Wikidepia/IndoT5-base-paraphrase"):
    """Load IndoT5 model"""
    print(f"🔧 Loading IndoT5: {model_name}")
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)
    print(f"✅ Model loaded on {device}")
    return tokenizer, model, device

def paraphrase_with_indot5(text, tokenizer, model, device, temperature=1.0, num_beams=10):
    """
    High-quality paraphrase dengan IndoT5
    
    Args:
        temperature: Higher = more diverse (0.8-1.2)
        num_beams: More beams = better quality (5-10)
    """
    input_text = f"paraphrase: {text}"
    inputs = tokenizer(
        input_text,
        return_tensors="pt",
        max_length=512,
        truncation=True,
        padding=True
    ).to(device)
    
    with torch.no_grad():
        outputs = model.generate(
            inputs.input_ids,
            max_length=512,
            num_beams=num_beams,           # More beams for quality
            temperature=temperature,        # Diversity
            do_sample=True,                # Enable sampling
            top_k=50,                      # Top-k sampling
            top_p=0.95,                    # Nucleus sampling
            early_stopping=True,
            no_repeat_ngram_size=3,        # Avoid repetition
            num_return_sequences=1
        )
    
    paraphrased = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return paraphrased

def match_flagged_with_docx(flagged_json, docx_path, threshold=0.5):
    """Match flagged texts with DOCX"""
    with open(flagged_json, 'r', encoding='utf-8') as f:
        flagged_texts = json.load(f)
    
    doc = docx.Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    print(f"\n🔍 Matching {len(flagged_texts)} flagged with {len(paragraphs)} paragraphs...")
    
    matches = []
    for i, flagged in enumerate(flagged_texts, 1):
        flagged_text = flagged['text']
        best_match = None
        best_ratio = 0
        best_idx = -1
        
        for idx, para in enumerate(paragraphs):
            ratio = SequenceMatcher(None, flagged_text.lower(), para.lower()).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = para
                best_idx = idx
        
        if best_ratio < 0.5:
            keywords = [w.lower() for w in flagged_text.split() if len(w) > 4]
            for idx, para in enumerate(paragraphs):
                keyword_matches = sum(1 for kw in keywords if kw in para.lower())
                keyword_ratio = keyword_matches / len(keywords) if keywords else 0
                if keyword_ratio > best_ratio:
                    best_ratio = keyword_ratio
                    best_match = para
                    best_idx = idx
        
        if best_ratio >= threshold:
            matches.append({
                "flagged_text": flagged_text,
                "matched_para": best_match,
                "similarity": round(best_ratio * 100, 2),
                "para_index": best_idx,
                "page": flagged['page']
            })
    
    print(f"✅ Matched: {len(matches)}/{len(flagged_texts)}")
    return matches

def paraphrase_all_clean(matches):
    """Pure IndoT5 paraphrase - NO tricks"""
    print(f"\n🎨 Paraphrasing {len(matches)} texts with IndoT5...")
    print("Strategy: High-quality natural paraphrase ONLY")
    print("="*70)
    
    tokenizer, model, device = load_indot5_model()
    results = []
    
    for match in tqdm(matches, desc="Paraphrasing"):
        original = match['matched_para']
        
        # Split long texts into sentences
        if len(original) > 400:
            sentences = [s.strip() + '.' for s in original.split('.') if s.strip()]
            paraphrased_parts = []
            
            for sent in sentences:
                if len(sent) > 10:
                    para = paraphrase_with_indot5(sent, tokenizer, model, device, 
                                                   temperature=1.0, num_beams=10)
                    paraphrased_parts.append(para)
                else:
                    paraphrased_parts.append(sent)
            
            paraphrased = ' '.join(paraphrased_parts)
        else:
            paraphrased = paraphrase_with_indot5(original, tokenizer, model, device,
                                                  temperature=1.0, num_beams=10)
        
        # NO invisible chars, NO unicode substitution - CLEAN OUTPUT
        results.append({
            "original": original,
            "paraphrased": paraphrased,  # Pure IndoT5 output
            "method": "indot5_clean",
            "page": match['page'],
            "similarity_before": match['similarity'],
            "para_index": match['para_index']
        })
    
    return results

def apply_to_docx(paraphrased_data, input_docx, output_docx):
    """Apply clean paraphrased texts to DOCX"""
    print(f"\n📄 Loading: {input_docx}")
    doc = docx.Document(input_docx)
    
    replaced = 0
    for item in tqdm(paraphrased_data, desc="Applying"):
        original = item['original']
        paraphrased = item['paraphrased']
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            
            ratio = SequenceMatcher(None, original.lower(), para.text.lower()).ratio()
            if ratio >= 0.7:
                # Replace keeping formatting
                if para.runs:
                    for run in para.runs:
                        run.text = ''
                    para.runs[0].text = paraphrased
                else:
                    para.text = paraphrased
                replaced += 1
                break
    
    print(f"✅ Replaced: {replaced}/{len(paraphrased_data)}")
    doc.save(output_docx)
    print(f"💾 Saved: {output_docx}")
    return replaced

def main():
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python clean_paraphrase.py <flagged.json> <original.docx>")
        sys.exit(1)
    
    flagged_json = sys.argv[1]
    docx_file = sys.argv[2]
    output_file = "testing_clean_bypassed.docx"
    
    print("="*70)
    print("🎓 CLEAN APPROACH: Pure IndoT5 Paraphrase")
    print("="*70)
    print("Strategy:")
    print("  ✅ High-quality IndoT5 paraphrase")
    print("  ✅ Natural language output")
    print("  ✅ CLEAN text layer (no tricks)")
    print("  ❌ NO invisible characters")
    print("  ❌ NO unicode substitution")
    print("="*70)
    
    # Step 1: Match
    matches = match_flagged_with_docx(flagged_json, docx_file, threshold=0.5)
    
    if not matches:
        print("❌ No matches found")
        return
    
    # Save matches
    with open("clean_matches.json", 'w', encoding='utf-8') as f:
        json.dump(matches, f, ensure_ascii=False, indent=2)
    
    # Step 2: Pure paraphrase
    paraphrased = paraphrase_all_clean(matches)
    
    # Save paraphrased
    with open("clean_paraphrased.json", 'w', encoding='utf-8') as f:
        json.dump(paraphrased, f, ensure_ascii=False, indent=2)
    
    # Step 3: Apply to DOCX
    apply_to_docx(paraphrased, docx_file, output_file)
    
    print("\n" + "="*70)
    print("✅ COMPLETE!")
    print("="*70)
    print(f"📄 Output: {output_file}")
    print(f"📊 Paraphrased: {len(paraphrased)} texts")
    print()
    print("📋 What's different from before:")
    print("  - NO invisible characters (clean OCR)")
    print("  - NO unicode substitution (clean text)")
    print("  - ONLY natural IndoT5 paraphrase")
    print("  - Should have CLEAN text layer")
    print()
    print("🎯 Next: Upload to Turnitin and compare!")
    print("="*70)

if __name__ == "__main__":
    main()
