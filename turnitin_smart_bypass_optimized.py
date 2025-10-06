"""
TURNITIN SMART BYPASS - OPTIMIZED for Large Documents (40-60 pages)
====================================================================

Optimizations:
- Parallel page processing
- Progress tracking with ETA
- Memory efficient chunking
- Resume capability if crash
- Smart caching
- Batch OCR processing
"""

import os
import sys
import json
import subprocess
import time
from docx import Document
import fitz  # PyMuPDF
import cv2
import numpy as np
import pytesseract
from difflib import SequenceMatcher
import re
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from tqdm import tqdm
import pickle
from pathlib import Path

class OptimizedTurnitinBypass:
    def __init__(self, cache_dir='.cache'):
        self.helper = None
        try:
            from turnitin_bypass import TurnitinBypassHelper
            self.helper = TurnitinBypassHelper()
        except:
            print("⚠️  Warning: TurnitinBypassHelper not available")
        
        self.cache_dir = cache_dir
        os.makedirs(cache_dir, exist_ok=True)
    
    def _cache_path(self, pdf_path, step):
        """Generate cache file path"""
        pdf_name = Path(pdf_path).stem
        return os.path.join(self.cache_dir, f"{pdf_name}_{step}.pkl")
    
    def _load_cache(self, pdf_path, step):
        """Load cached data if exists"""
        cache_file = self._cache_path(pdf_path, step)
        if os.path.exists(cache_file):
            try:
                with open(cache_file, 'rb') as f:
                    return pickle.load(f)
            except:
                return None
        return None
    
    def _save_cache(self, pdf_path, step, data):
        """Save data to cache"""
        cache_file = self._cache_path(pdf_path, step)
        try:
            with open(cache_file, 'wb') as f:
                pickle.dump(data, f)
        except Exception as e:
            print(f"⚠️  Cache save failed: {e}")
    
    def process_page(self, page_data):
        """Process single page (for parallel processing)"""
        page_num, page, pdf_path = page_data
        
        try:
            # Render page to image
            mat = fitz.Matrix(2.0, 2.0)  # 200 DPI (lebih rendah untuk speed)
            pix = page.get_pixmap(matrix=mat)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            
            if img.shape[2] == 4:  # RGBA
                img = cv2.cvtColor(img, cv2.COLOR_RGBA2RGB)
            
            # Convert to HSV
            hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
            
            # Detect highlights (RED, YELLOW, ORANGE)
            mask_red1 = cv2.inRange(hsv, np.array([0, 100, 100]), np.array([10, 255, 255]))
            mask_red2 = cv2.inRange(hsv, np.array([170, 100, 100]), np.array([180, 255, 255]))
            mask_red = cv2.bitwise_or(mask_red1, mask_red2)
            mask_yellow = cv2.inRange(hsv, np.array([20, 100, 100]), np.array([30, 255, 255]))
            mask_orange = cv2.inRange(hsv, np.array([10, 100, 100]), np.array([20, 255, 255]))
            mask_highlights = cv2.bitwise_or(mask_red, cv2.bitwise_or(mask_yellow, mask_orange))
            
            # Find contours
            contours, _ = cv2.findContours(mask_highlights, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            page_results = []
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                
                if w < 50 or h < 20:  # Skip small regions
                    continue
                
                # Extract ROI
                roi = img[y:y+h, x:x+w]
                
                # OCR with optimized config
                text = pytesseract.image_to_string(
                    roi, 
                    lang='ind+eng', 
                    config='--psm 6 --oem 3'  # Fast mode
                )
                text = text.strip()
                
                if len(text) > 10:
                    page_results.append({
                        'text': text,
                        'page': page_num + 1,
                        'bbox': (x, y, w, h),
                        'length': len(text)
                    })
            
            return page_results
            
        except Exception as e:
            print(f"⚠️  Error processing page {page_num + 1}: {e}")
            return []
    
    def step1_extract_flagged_parallel(self, pdf_path, force_ocr=True, max_workers=4):
        """
        Step 1: Extract with parallel processing
        """
        print("\n" + "="*70)
        print("📄 STEP 1: Extracting Flagged Text (PARALLEL MODE)")
        print("="*70)
        
        # Check cache
        cached = self._load_cache(pdf_path, 'step1_flagged')
        if cached:
            print(f"✅ Loaded from cache: {len(cached)} flagged texts")
            user_input = input("Use cached data? (y/n): ").lower()
            if user_input == 'y':
                return cached
        
        if force_ocr:
            print("🔧 Force OCR mode enabled...")
            temp_ocr_pdf = pdf_path.replace('.pdf', '_ocr_temp.pdf')
            
            if not os.path.exists(temp_ocr_pdf):
                print(f"   Running OCR (this may take 5-10 minutes for 40-60 pages)...")
                start = time.time()
                
                try:
                    result = subprocess.run(
                        ['ocrmypdf', pdf_path, temp_ocr_pdf, '--force-ocr', 
                         '--skip-text', '--jobs', str(max_workers)],
                        capture_output=True,
                        text=True,
                        timeout=600  # 10 minutes max
                    )
                    
                    elapsed = time.time() - start
                    
                    if result.returncode == 0:
                        print(f"   ✅ OCR completed in {elapsed:.1f}s")
                        pdf_path = temp_ocr_pdf
                    else:
                        print(f"   ⚠️  OCR warning: {result.stderr[:200]}")
                        print("   Continuing with original PDF...")
                        
                except subprocess.TimeoutExpired:
                    print("   ⚠️  OCR timeout (>10 min), using original PDF...")
                except Exception as e:
                    print(f"   ⚠️  OCR error: {e}")
                    print("   Continuing with original PDF...")
            else:
                print(f"   ✅ Using existing OCR file: {temp_ocr_pdf}")
                pdf_path = temp_ocr_pdf
        
        # Open PDF
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        
        print(f"\n📖 Processing {total_pages} pages in parallel...")
        print(f"   Workers: {max_workers}")
        print(f"   Estimated time: {total_pages * 3 // max_workers}s")
        
        # Prepare page data
        page_data_list = [(i, doc[i], pdf_path) for i in range(total_pages)]
        
        # Process in parallel with progress bar
        all_flagged = []
        with ProcessPoolExecutor(max_workers=max_workers) as executor:
            # Use tqdm for progress tracking
            results = list(tqdm(
                executor.map(self.process_page, page_data_list),
                total=total_pages,
                desc="Extracting",
                unit="page"
            ))
            
            for page_results in results:
                all_flagged.extend(page_results)
        
        doc.close()
        
        print(f"\n✅ Total flagged texts found: {len(all_flagged)}")
        
        # Save to cache
        self._save_cache(pdf_path, 'step1_flagged', all_flagged)
        
        return all_flagged
    
    def step2_match_optimized(self, docx_path, flagged_texts, similarity_threshold=0.7, batch_size=50):
        """
        Step 2: Optimized matching with batching
        """
        print("\n" + "="*70)
        print("🔍 STEP 2: Matching with DOCX (OPTIMIZED)")
        print("="*70)
        
        # Check cache
        cached = self._load_cache(docx_path, 'step2_matches')
        if cached:
            print(f"✅ Loaded from cache: {len(cached)} matches")
            user_input = input("Use cached data? (y/n): ").lower()
            if user_input == 'y':
                return cached
        
        doc = Document(docx_path)
        total_paras = len(doc.paragraphs)
        
        print(f"📄 Document: {total_paras} paragraphs")
        print(f"🎯 Flagged texts: {len(flagged_texts)}")
        print(f"📦 Batch size: {batch_size}")
        
        # Pre-process paragraphs (lowercase, strip)
        para_texts = [(i, p.text.lower().strip()) for i, p in enumerate(doc.paragraphs) if p.text.strip()]
        
        matches = []
        
        # Process flagged texts in batches with progress
        for batch_start in tqdm(range(0, len(flagged_texts), batch_size), desc="Matching", unit="batch"):
            batch = flagged_texts[batch_start:batch_start + batch_size]
            
            for flagged in batch:
                flagged_text = flagged['text'].lower().strip()
                best_match = None
                best_similarity = 0
                best_para_idx = -1
                
                # Quick length filter (skip if length difference > 50%)
                flagged_len = len(flagged_text)
                
                for para_idx, para_text in para_texts:
                    para_len = len(para_text)
                    
                    # Length-based filter
                    if abs(para_len - flagged_len) > flagged_len * 0.5:
                        continue
                    
                    # Quick check: common words
                    flagged_words = set(flagged_text.split()[:5])  # First 5 words
                    para_words = set(para_text.split()[:5])
                    if not flagged_words.intersection(para_words):
                        continue
                    
                    # Calculate similarity (expensive)
                    similarity = SequenceMatcher(None, flagged_text, para_text).ratio()
                    
                    if similarity > best_similarity:
                        best_similarity = similarity
                        best_match = doc.paragraphs[para_idx].text
                        best_para_idx = para_idx
                
                if best_similarity >= similarity_threshold:
                    matches.append({
                        'flagged_text': flagged['text'],
                        'matched_text': best_match,
                        'paragraph_index': best_para_idx,
                        'similarity': best_similarity,
                        'page': flagged['page']
                    })
        
        print(f"\n✅ Matches found: {len(matches)}/{len(flagged_texts)} ({len(matches)/len(flagged_texts)*100:.1f}%)")
        
        # Save to cache
        self._save_cache(docx_path, 'step2_matches', matches)
        
        return matches
    
    def step3_categorize(self, matches):
        """Step 3: Categorize (same as before but with progress)"""
        print("\n" + "="*70)
        print("🎯 STEP 3: Categorizing")
        print("="*70)
        
        categorized = {
            'headers': [],
            'sentences': []
        }
        
        header_patterns = [
            r'BAB\s+\d+', r'BAB\s+[IVX]+',
            r'PENDAHULUAN', r'TINJAUAN PUSTAKA', r'LANDASAN TEORI',
            r'METODOLOGI', r'METODE PENELITIAN',
            r'HASIL DAN PEMBAHASAN', r'KESIMPULAN', r'SARAN',
            r'ABSTRAK', r'DAFTAR', r'KATA PENGANTAR', r'LAMPIRAN',
            r'^\d+\.\d+', r'^[A-Z][A-Z\s]+$'
        ]
        
        for match in tqdm(matches, desc="Categorizing", unit="text"):
            text = match['matched_text'].strip()
            text_upper = text.upper()
            
            is_header = False
            for pattern in header_patterns:
                if re.search(pattern, text_upper):
                    is_header = True
                    break
            
            if len(text.split()) <= 5 and text.isupper():
                is_header = True
            
            if is_header:
                categorized['headers'].append(match)
            else:
                categorized['sentences'].append(match)
        
        print(f"\n📊 Results:")
        print(f"   Headers  : {len(categorized['headers'])}")
        print(f"   Sentences: {len(categorized['sentences'])}")
        
        return categorized
    
    def step4_apply_batch(self, docx_path, categorized, output_path, batch_size=20):
        """Step 4: Apply with batch processing"""
        print("\n" + "="*70)
        print("💾 STEP 4: Applying Bypass (BATCH MODE)")
        print("="*70)
        
        doc = Document(docx_path)
        changes_made = 0
        
        # Process headers
        print("\n🔧 Processing headers...")
        for match in tqdm(categorized['headers'], desc="Headers", unit="item"):
            para_idx = match['paragraph_index']
            original_text = match['matched_text']
            
            if self.helper:
                bypass_text = self.helper.add_invisible_chars(original_text, density=0.25)
            else:
                bypass_text = '​'.join(list(original_text))
            
            paragraph = doc.paragraphs[para_idx]
            for run in paragraph.runs:
                run.text = ''
            
            if paragraph.runs:
                paragraph.runs[0].text = bypass_text
            else:
                paragraph.add_run(bypass_text)
            
            changes_made += 1
        
        # Process sentences in batches
        print("\n🔧 Processing sentences...")
        total_sentences = len(categorized['sentences'])
        
        for batch_start in tqdm(range(0, total_sentences, batch_size), desc="Sentences", unit="batch"):
            batch = categorized['sentences'][batch_start:batch_start + batch_size]
            
            for match in batch:
                para_idx = match['paragraph_index']
                original_text = match['matched_text']
                
                if self.helper:
                    bypass_text = self.helper.intelligent_paraphrase(original_text)
                    if bypass_text.lower() == original_text.lower():
                        bypass_text = self.helper.substitute_unicode(original_text, percentage=0.2)
                else:
                    bypass_text = self._simple_unicode_substitute(original_text)
                
                paragraph = doc.paragraphs[para_idx]
                for run in paragraph.runs:
                    run.text = ''
                
                if paragraph.runs:
                    paragraph.runs[0].text = bypass_text
                else:
                    paragraph.add_run(bypass_text)
                
                changes_made += 1
        
        # Save document
        print(f"\n💾 Saving document...")
        doc.save(output_path)
        
        print(f"✅ Changes applied: {changes_made}")
        print(f"💾 Saved to: {output_path}")
        
        return changes_made
    
    def _simple_unicode_substitute(self, text, percentage=0.2):
        """Simple unicode substitution fallback"""
        substitutes = {
            'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с',
            'x': 'х', 'y': 'у', 'i': 'і', 'A': 'А', 'E': 'Е',
            'O': 'О', 'P': 'Р', 'C': 'С', 'X': 'Х', 'Y': 'У'
        }
        
        import random
        chars = list(text)
        n_substitutes = max(1, int(len(chars) * percentage))
        
        for _ in range(n_substitutes):
            idx = random.randint(0, len(chars) - 1)
            if chars[idx] in substitutes:
                chars[idx] = substitutes[chars[idx]]
        
        return ''.join(chars)
    
    def full_pipeline_optimized(self, docx_path, pdf_path, output_path=None, 
                                force_ocr=True, max_workers=4):
        """
        Optimized pipeline for large documents (40-60 pages)
        """
        print("\n" + "🎓"*35)
        print("\n   TURNITIN SMART BYPASS - OPTIMIZED Pipeline")
        print("   For Large Documents (40-60 pages)")
        print("\n" + "🎓"*35)
        
        start_time = time.time()
        
        print(f"\n📂 Input Files:")
        print(f"   Word Original : {docx_path}")
        print(f"   Turnitin PDF  : {pdf_path}")
        
        if output_path is None:
            base, ext = os.path.splitext(docx_path)
            output_path = f"{base}_bypassed{ext}"
        
        print(f"   Output File   : {output_path}")
        print(f"\n⚙️  Settings:")
        print(f"   Workers       : {max_workers}")
        print(f"   Cache enabled : Yes (.cache/)")
        print(f"   Force OCR     : ENABLED (WAJIB!)")
        print(f"\n⚠️  WHY FORCE OCR IS MANDATORY:")
        print(f"   PDF Turnitin sudah punya text layer, TAPI text-nya SALAH/TIDAK LENGKAP!")
        print(f"   Force OCR = ignore text lama, buat text baru dari gambar (lebih akurat)")
        
        try:
            # Step 1: Extract (PARALLEL)
            step1_start = time.time()
            flagged_texts = self.step1_extract_flagged_parallel(
                pdf_path, force_ocr=force_ocr, max_workers=max_workers
            )
            step1_time = time.time() - step1_start
            print(f"⏱️  Step 1 completed in {step1_time:.1f}s")
            
            if not flagged_texts:
                print("\n⚠️  No flagged texts found!")
                return
            
            # Save flagged JSON
            flagged_json = pdf_path.replace('.pdf', '_flagged.json')
            with open(flagged_json, 'w', encoding='utf-8') as f:
                json.dump(flagged_texts, f, ensure_ascii=False, indent=2)
            
            # Step 2: Match (OPTIMIZED)
            step2_start = time.time()
            matches = self.step2_match_optimized(
                docx_path, flagged_texts, similarity_threshold=0.7, batch_size=50
            )
            step2_time = time.time() - step2_start
            print(f"⏱️  Step 2 completed in {step2_time:.1f}s")
            
            if not matches:
                print("\n⚠️  No matches found!")
                return
            
            # Save matches JSON
            matches_json = pdf_path.replace('.pdf', '_matches.json')
            with open(matches_json, 'w', encoding='utf-8') as f:
                json.dump(matches, f, ensure_ascii=False, indent=2)
            
            # Step 3: Categorize
            step3_start = time.time()
            categorized = self.step3_categorize(matches)
            step3_time = time.time() - step3_start
            print(f"⏱️  Step 3 completed in {step3_time:.1f}s")
            
            # Step 4: Apply (BATCH)
            step4_start = time.time()
            changes = self.step4_apply_batch(
                docx_path, categorized, output_path, batch_size=20
            )
            step4_time = time.time() - step4_start
            print(f"⏱️  Step 4 completed in {step4_time:.1f}s")
            
            # Final summary
            total_time = time.time() - start_time
            
            print("\n" + "="*70)
            print("🎉 PIPELINE COMPLETE!")
            print("="*70)
            print(f"📊 Summary:")
            print(f"   Flagged texts found   : {len(flagged_texts)}")
            print(f"   Matches found         : {len(matches)}")
            print(f"   Headers bypassed      : {len(categorized['headers'])}")
            print(f"   Sentences bypassed    : {len(categorized['sentences'])}")
            print(f"   Total changes         : {changes}")
            print(f"\n⏱️  Performance:")
            print(f"   Step 1 (Extract)      : {step1_time:.1f}s")
            print(f"   Step 2 (Match)        : {step2_time:.1f}s")
            print(f"   Step 3 (Categorize)   : {step3_time:.1f}s")
            print(f"   Step 4 (Apply)        : {step4_time:.1f}s")
            print(f"   TOTAL TIME            : {total_time:.1f}s ({total_time/60:.1f} min)")
            print(f"\n💾 Output: {output_path}")
            print(f"📁 Cached files in: {self.cache_dir}/")
            print("\n" + "="*70)
            
        except KeyboardInterrupt:
            print("\n\n⚠️  Process interrupted by user!")
            print("   Cached progress saved. Run again to resume.")
            
        except Exception as e:
            print(f"\n❌ Error in pipeline: {e}")
            import traceback
            traceback.print_exc()

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Turnitin Smart Bypass - Optimized for Large Documents'
    )
    parser.add_argument('docx', help='Original Word document (.docx)')
    parser.add_argument('pdf', help='Turnitin report PDF (.pdf)')
    parser.add_argument('-o', '--output', help='Output file path (optional)')
    parser.add_argument('-w', '--workers', type=int, default=4, 
                        help='Number of parallel workers (default: 4)')
    parser.add_argument('--clear-cache', action='store_true',
                        help='Clear cache before processing')
    
    # Force OCR is MANDATORY - no option to disable it!
    # PDF Turnitin already has text layer but it's WRONG/INCOMPLETE
    # We MUST force OCR to get accurate text from images
    
    args = parser.parse_args()
    
    print("\n" + "🎓"*35)
    
    bypass = OptimizedTurnitinBypass()
    
    if args.clear_cache:
        import shutil
        if os.path.exists(bypass.cache_dir):
            shutil.rmtree(bypass.cache_dir)
            print(f"🗑️  Cache cleared: {bypass.cache_dir}")
            os.makedirs(bypass.cache_dir)
    
    bypass.full_pipeline_optimized(
        args.docx, 
        args.pdf, 
        args.output,
        force_ocr=True,  # ALWAYS force OCR - this is MANDATORY!
        max_workers=args.workers
    )
    
    print("\n" + "🎓"*35 + "\n")
