#!/usr/bin/env python3
"""
Apply paraphrased results ke DOCX dan generate final file
"""

import json
import docx
from pathlib import Path
from difflib import SequenceMatcher

def find_and_replace_paragraph(doc, original_text, new_text, similarity_threshold=0.7):
    """
    Find paragraph yang match dengan original_text dan replace dengan new_text
    Returns: True if replaced, False otherwise
    """
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        
        # Check similarity
        ratio = SequenceMatcher(None, original_text.lower(), para.text.lower()).ratio()
        
        if ratio >= similarity_threshold:
            # Found match - replace
            # Keep original formatting
            if para.runs:
                # Clear existing runs
                for run in para.runs:
                    run.text = ''
                # Add new text to first run
                para.runs[0].text = new_text
            else:
                # No runs, set directly
                para.text = new_text
            
            return True
    
    return False

def apply_paraphrased_to_docx(paraphrased_json, input_docx, output_docx):
    """
    Apply semua paraphrased texts ke DOCX dan save
    """
    # Load paraphrased data
    with open(paraphrased_json, 'r', encoding='utf-8') as f:
        paraphrased_data = json.load(f)
    
    print(f"\n📄 Loading DOCX: {input_docx}")
    doc = docx.Document(input_docx)
    
    print(f"📊 Total paragraphs in DOCX: {len(doc.paragraphs)}")
    print(f"📊 Total paraphrased texts: {len(paraphrased_data)}")
    
    print(f"\n🔧 Applying changes...")
    print("="*70)
    
    replaced_count = 0
    failed_count = 0
    
    # Sort by para_index untuk konsistensi
    paraphrased_data_sorted = sorted(paraphrased_data, key=lambda x: x.get('para_index', 0))
    
    for i, item in enumerate(paraphrased_data_sorted, 1):
        original = item['original']
        paraphrased = item['paraphrased']
        method = item.get('method', 'unknown')
        page = item.get('page', '?')
        
        # Try to replace
        success = find_and_replace_paragraph(doc, original, paraphrased, similarity_threshold=0.7)
        
        if success:
            replaced_count += 1
            method_icon = "🔤" if method == "invisible_chars" else "🎨"
            print(f"   ✅ {method_icon} {i}/{len(paraphrased_data)} - Page {page}: {original[:50]}...")
        else:
            failed_count += 1
            print(f"   ❌ {i}/{len(paraphrased_data)} - Failed to find: {original[:50]}...")
    
    print("="*70)
    print(f"\n📊 Results:")
    print(f"   ✅ Replaced : {replaced_count}/{len(paraphrased_data)}")
    print(f"   ❌ Failed   : {failed_count}/{len(paraphrased_data)}")
    
    # Save output
    print(f"\n💾 Saving to: {output_docx}")
    doc.save(output_docx)
    
    print(f"✅ Done! File saved successfully.")
    
    return replaced_count, failed_count

def main():
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python apply_to_docx.py <paraphrased.json> <input.docx> [output.docx]")
        print("\nExample:")
        print("  python apply_to_docx.py testing_paraphrased.json testing.docx testing_bypassed.docx")
        sys.exit(1)
    
    paraphrased_json = sys.argv[1]
    input_docx = sys.argv[2]
    
    # Auto-generate output name if not provided
    if len(sys.argv) > 3:
        output_docx = sys.argv[3]
    else:
        input_path = Path(input_docx)
        output_docx = str(input_path.parent / f"{input_path.stem}_bypassed{input_path.suffix}")
    
    print("="*70)
    print("🎓 TURNITIN BYPASS - Apply to DOCX")
    print("="*70)
    print(f"📥 Input JSON : {paraphrased_json}")
    print(f"📥 Input DOCX : {input_docx}")
    print(f"📤 Output DOCX: {output_docx}")
    print("="*70)
    
    replaced, failed = apply_paraphrased_to_docx(paraphrased_json, input_docx, output_docx)
    
    print("\n" + "="*70)
    print("🎉 COMPLETE!")
    print("="*70)
    print(f"✅ Successfully replaced: {replaced} paragraphs")
    if failed > 0:
        print(f"⚠️  Failed to replace: {failed} paragraphs")
    print(f"\n📄 Output file: {output_docx}")
    print("="*70)
    print("\n💡 Next steps:")
    print("   1. Open the bypassed file in Word")
    print("   2. Review the changes")
    print("   3. Upload to Turnitin for checking")
    print("="*70)

if __name__ == "__main__":
    main()
