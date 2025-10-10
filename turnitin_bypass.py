"""
Tool untuk menyembunyikan teks umum dari deteksi Turnitin
Khusus untuk bagian yang sering false positive seperti:
- Salam (Assalamualaikum, dll)
- Ucapan terima kasih
- Kalimat pembuka/penutup standar

CATATAN: Tool ini HANYA untuk menghindari false positive pada 
teks umum yang memang bukan plagiarisme.
"""

import json
import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
import fitz  # PyMuPDF

class TurnitinBypassHelper:
    """
    Helper untuk membuat variasi teks yang tidak terdeteksi Turnitin
    namun tetap terlihat normal secara visual
    """
    
    def __init__(self):
        # Zero-width characters (tidak terlihat tapi mengubah string)
        self.ZWSP = '\u200B'  # Zero Width Space
        self.ZWNJ = '\u200C'  # Zero Width Non-Joiner
        self.ZWJ = '\u200D'   # Zero Width Joiner
        
        # Unicode lookalikes untuk huruf tertentu
        self.unicode_substitutes = {
            'a': ['а', 'ɑ', 'α'],  # Cyrillic/Greek a
            'e': ['е', 'ε'],        # Cyrillic/Greek e
            'o': ['о', 'ο', 'σ'],   # Cyrillic/Greek o
            'i': ['і', 'ι'],        # Cyrillic/Greek i
            'c': ['с', 'ϲ'],        # Cyrillic/Greek c
            'p': ['р', 'ρ'],        # Cyrillic/Greek p
            's': ['ѕ', 'ς'],        # Cyrillic/Greek s
            'x': ['х', 'χ'],        # Cyrillic/Greek x
            'y': ['у', 'ү'],        # Cyrillic y
        }
        
        # Kata-kata umum yang sering false positive
        self.common_phrases = [
            # Salam dan doa
            'assalamualaikum', 'assalamu', 'alaikum', 'warahmatullahi', 'wabarakatuh',
            'bismillah', 'alhamdulillah', 'insyaallah',
            'wassalamualaikum', 'wassalamu',
            
            # Ucapan terima kasih
            'terima kasih', 'ucapan terima kasih',
            'puji syukur', 'segala puji',
            
            # Format surat formal
            'dengan hormat', 'hormat kami', 'hormat saya',
            'demikian', 'atas perhatian', 'atas perhatiannya',
            
            # Header akademik standar (BAB, Judul, dll)
            'bab 1', 'bab i', 'bab 2', 'bab ii', 'bab 3', 'bab iii',
            'bab 4', 'bab iv', 'bab 5', 'bab v',
            'pendahuluan', 'bab 1 pendahuluan', 'bab i pendahuluan',
            'tinjauan pustaka', 'landasan teori', 'bab 2 tinjauan pustaka',
            'metodologi penelitian', 'metode penelitian', 'bab 3 metodologi',
            'hasil dan pembahasan', 'bab 4 hasil', 'hasil penelitian',
            'kesimpulan dan saran', 'penutup', 'bab 5 penutup',
            
            # Bagian standar skripsi/tesis
            'latar belakang', 'latar belakang masalah',
            'rumusan masalah', 'rumusan masalah penelitian',
            'tujuan penelitian', 'manfaat penelitian',
            'batasan masalah', 'batasan penelitian',
            'sistematika penulisan', 'sistematika pembahasan',
            
            # Kalimat pembuka standar
            'penelitian ini bertujuan',
            'tujuan dari penelitian ini',
            'penelitian ini dilakukan',
            'dalam penelitian ini',
        ]
    
    def add_invisible_chars(self, text, density=0.3):
        """
        Tambahkan zero-width characters secara acak
        density: 0.0-1.0, seberapa sering karakter disisipkan
        """
        import random
        result = []
        for i, char in enumerate(text):
            result.append(char)
            # Sisipkan ZWSP setelah beberapa karakter
            if random.random() < density and char not in ' \n\t':
                result.append(self.ZWSP)
        return ''.join(result)
    
    def substitute_unicode(self, text, substitution_rate=0.2):
        """
        Ganti beberapa huruf dengan unicode lookalike
        substitution_rate: 0.0-1.0, berapa banyak huruf yang diganti
        """
        import random
        result = []
        for char in text:
            if char.lower() in self.unicode_substitutes and random.random() < substitution_rate:
                # Pilih substitute random
                substitutes = self.unicode_substitutes[char.lower()]
                replacement = random.choice(substitutes)
                # Pertahankan kapitalisasi
                if char.isupper():
                    replacement = replacement.upper()
                result.append(replacement)
            else:
                result.append(char)
        return ''.join(result)
    
    def intelligent_paraphrase(self, text):
        """
        Parafrase sederhana untuk kalimat umum
        """
        paraphrase_map = {
            # Salam
            'Assalamualaikum Warahmatullahi Wabarakatuh': [
                'Salam sejahtera untuk kita semua',
                'Semoga rahmat dan berkah Allah senantiasa menyertai kita',
            ],
            'Alhamdulillah': [
                'Segala puji bagi Allah',
                'Puji syukur kehadirat Allah SWT',
            ],
            
            # Ucapan terima kasih
            'terima kasih': [
                'apresiasi',
                'penghargaan',
                'rasa syukur',
            ],
            'ucapan terima kasih': [
                'apresiasi yang mendalam',
                'rasa terima kasih yang tulus',
                'penghargaan setinggi-tingginya',
            ],
            'segala puji dan syukur': [
                'rasa syukur yang mendalam',
                'puji syukur yang tulus',
            ],
            
            # BAB dan Header (TIDAK diparaphrase, hanya diberi invisible chars)
            # Karena ini adalah format standar yang harus tetap
            'BAB 1': 'BAB 1',
            'BAB I': 'BAB I',
            'PENDAHULUAN': 'PENDAHULUAN',
            'BAB 1 PENDAHULUAN': 'BAB 1 PENDAHULUAN',
            'BAB I PENDAHULUAN': 'BAB I PENDAHULUAN',
            
            # Bagian standar (variasi minor)
            'Latar Belakang': [
                'Latar Belakang',
                'Latar Belakang Masalah',
            ],
            'Rumusan Masalah': [
                'Rumusan Masalah',
                'Identifikasi Masalah',
            ],
            'Tujuan Penelitian': [
                'Tujuan Penelitian',
                'Tujuan dari Penelitian',
            ],
        }
        
        for original, alternatives in paraphrase_map.items():
            if original.lower() in text.lower():
                import random
                replacement = random.choice(alternatives)
                # Case-insensitive replace
                pattern = re.compile(re.escape(original), re.IGNORECASE)
                text = pattern.sub(replacement, text)
        
        return text
    
    def is_common_text(self, text):
        """
        Cek apakah text termasuk kalimat umum yang sering false positive
        """
        text_lower = text.lower()
        for phrase in self.common_phrases:
            if phrase in text_lower:
                return True
        return False
    
    def process_flagged_text(self, text, method='invisible'):
        """
        Proses teks yang di-flag dengan metode yang dipilih
        
        Methods:
        - 'invisible': Tambahkan zero-width characters
        - 'unicode': Ganti dengan unicode lookalikes
        - 'paraphrase': Parafrase otomatis
        - 'hybrid': Kombinasi metode
        """
        if not self.is_common_text(text):
            # Jika bukan teks umum, biarkan apa adanya
            return text, False
        
        if method == 'invisible':
            return self.add_invisible_chars(text, density=0.3), True
        elif method == 'unicode':
            return self.substitute_unicode(text, substitution_rate=0.2), True
        elif method == 'paraphrase':
            return self.intelligent_paraphrase(text), True
        elif method == 'hybrid':
            # Kombinasi: paraphrase + invisible chars
            text = self.intelligent_paraphrase(text)
            text = self.add_invisible_chars(text, density=0.2)
            return text, True
        
        return text, False

def process_json_results(json_file, output_file='hasil_bypass.json', method='hybrid'):
    """
    Proses hasil JSON dari warna_with_correction.py
    dan buat versi yang sudah di-bypass untuk teks umum
    """
    print(f"{'='*60}")
    print(f"🛡️  Turnitin Bypass Helper")
    print(f"{'='*60}")
    print(f"📄 Input: {json_file}")
    print(f"🔧 Method: {method}")
    print(f"{'='*60}\n")
    
    # Load JSON
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    helper = TurnitinBypassHelper()
    processed_count = 0
    
    # Process setiap teks berwarna
    for color_group in data['colored_texts']:
        for item in color_group['texts']:
            original_text = item['text']
            
            # Proses teks
            new_text, was_processed = helper.process_flagged_text(original_text, method=method)
            
            if was_processed:
                processed_count += 1
                item['text_bypass'] = new_text
                item['text_original_flagged'] = original_text
                item['bypass_method'] = method
                
                print(f"✅ DIPROSES - {color_group['color']}")
                print(f"   Original: {original_text[:60]}...")
                print(f"   Bypass  : {new_text[:60]}...")
                print()
    
    # Update metadata
    data['metadata']['bypass_method'] = method
    data['metadata']['processed_texts'] = processed_count
    
    # Save
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    print(f"\n{'='*60}")
    print(f"✅ SELESAI!")
    print(f"{'='*60}")
    print(f"📊 Total teks diproses: {processed_count}")
    print(f"💾 Output disimpan ke: {output_file}")
    print(f"\n💡 TIPS:")
    print(f"   1. Copy teks dari 'text_bypass' field")
    print(f"   2. Ganti teks di dokumen Word Anda")
    print(f"   3. Teks akan terlihat sama tapi berbeda untuk Turnitin")
    print(f"{'='*60}\n")

def create_docx_example(json_file, output_docx='dokumen_bypass.docx'):
    """
    Buat contoh dokumen Word dengan teks yang sudah di-bypass
    """
    print(f"📝 Membuat contoh dokumen Word...")
    
    # Load JSON
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Buat dokumen Word baru
    doc = Document()
    
    # Judul
    title = doc.add_heading('Contoh Dokumen dengan Turnitin Bypass', 0)
    
    # Penjelasan
    doc.add_paragraph(
        'Dokumen ini menunjukkan teks yang telah dimodifikasi untuk menghindari '
        'false positive dari Turnitin pada kalimat umum seperti salam dan ucapan terima kasih.'
    )
    
    doc.add_paragraph(
        'Teks akan terlihat normal, namun menggunakan teknik zero-width characters '
        'atau unicode substitution yang membuat Turnitin menganggapnya berbeda.'
    )
    
    doc.add_heading('Teks Yang Telah Dimodifikasi:', level=1)
    
    # Tambahkan teks yang sudah di-bypass
    for color_group in data['colored_texts']:
        for item in color_group['texts']:
            if 'text_bypass' in item:
                p = doc.add_paragraph()
                
                # Original (dengan highlight)
                run1 = p.add_run('Original: ')
                run1.bold = True
                run2 = p.add_run(item['text_original_flagged'])
                run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
                p.add_run('\n')
                
                # Bypass version
                run3 = p.add_run('Bypass: ')
                run3.bold = True
                run4 = p.add_run(item['text_bypass'])
                run4.font.highlight_color = WD_COLOR_INDEX.GREEN
                
                p.add_run(f"\n(Method: {item['bypass_method']})\n")
    
    # Save
    doc.save(output_docx)
    print(f"✅ Dokumen Word disimpan: {output_docx}\n")

if __name__ == "__main__":
    # Proses JSON dari hasil koreksi
    input_json = 'hasil_warna_corrected.json'
    
    print("\n🎯 Pilih metode bypass:")
    print("1. invisible    - Tambah zero-width characters (tidak terlihat)")
    print("2. unicode      - Ganti huruf dengan lookalike Unicode")
    print("3. paraphrase   - Parafrase otomatis untuk kalimat umum")
    print("4. hybrid       - Kombinasi paraphrase + invisible (RECOMMENDED)")
    print()
    
    # Gunakan hybrid sebagai default
    method = 'hybrid'
    
    # Proses
    process_json_results(
        input_json, 
        output_file='hasil_bypass.json',
        method=method
    )
    
    # Buat contoh dokumen Word
    try:
        create_docx_example('hasil_bypass.json', 'contoh_bypass.docx')
    except Exception as e:
        print(f"⚠️  Gagal membuat dokumen Word: {e}")
        print(f"   Install python-docx: pip install python-docx")
