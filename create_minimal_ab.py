#!/usr/bin/env python3
"""
Extract ONLY specific content from flagged pages
Creates minimal DOCX files for quick Turnitin testing
"""

import json
import docx
from docx import Document
from docx.shared import Pt
from difflib import SequenceMatcher

def create_minimal_test_files(original_docx, flagged_json, paraphrased_json, target_pages):
    """
    Create 2 minimal files containing ONLY the flagged texts from target pages
    File A: Original texts
    File B: Paraphrased texts
    """
    
    print("="*70)
    print("📄 CREATING MINIMAL A/B TEST FILES")
    print("="*70)
    print(f"Target pages: {target_pages}")
    
    # Load data
    with open(flagged_json, 'r', encoding='utf-8') as f:
        all_flagged = json.load(f)
    
    with open(paraphrased_json, 'r', encoding='utf-8') as f:
        paraphrased = json.load(f)
    
    # Filter by target pages
    target_flagged = [f for f in all_flagged if f['page'] in target_pages]
    target_para = [p for p in paraphrased if p['page'] in target_pages]
    
    print(f"\n📊 Filtered data:")
    print(f"   Flagged texts: {len(target_flagged)}")
    print(f"   Paraphrased: {len(target_para)}")
    
    # Load original DOCX to get actual paragraphs
    doc = Document(original_docx)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Match and extract
    original_texts = []
    modified_texts = []
    
    for para_item in target_para:
        original = para_item['original']
        paraphrased_text = para_item['paraphrased']
        page = para_item['page']
        
        original_texts.append({
            'text': original,
            'page': page
        })
        modified_texts.append({
            'text': paraphrased_text,
            'page': page
        })
    
    # Create File A: Original
    output_a = f"ORIGINAL_excerpt_pages_{'_'.join(map(str, target_pages))}.docx"
    doc_a = Document()
    doc_a.add_heading(f'Original Texts - Pages {", ".join(map(str, target_pages))}', 0)
    
    for i, item in enumerate(original_texts, 1):
        doc_a.add_heading(f'Text {i} (Page {item["page"]})', level=2)
        doc_a.add_paragraph(item['text'])
        doc_a.add_paragraph()  # Space
    
    doc_a.save(output_a)
    print(f"\n✅ File A created: {output_a}")
    print(f"   Contains: {len(original_texts)} original texts")
    
    # Create File B: Modified
    output_b = f"MODIFIED_excerpt_pages_{'_'.join(map(str, target_pages))}.docx"
    doc_b = Document()
    doc_b.add_heading(f'Modified Texts - Pages {", ".join(map(str, target_pages))}', 0)
    
    for i, item in enumerate(modified_texts, 1):
        doc_b.add_heading(f'Text {i} (Page {item["page"]})', level=2)
        doc_b.add_paragraph(item['text'])
        doc_b.add_paragraph()  # Space
    
    doc_b.save(output_b)
    print(f"✅ File B created: {output_b}")
    print(f"   Contains: {len(modified_texts)} paraphrased texts")
    
    print("\n" + "="*70)
    print("🎯 A/B TESTING GUIDE")
    print("="*70)
    print(f"📤 Upload BOTH files to Turnitin:")
    print(f"   1. {output_a}")
    print(f"   2. {output_b}")
    print()
    print("📊 Expected results:")
    print("   - File A (original) = Higher similarity")
    print("   - File B (modified) = LOWER similarity ✅")
    print()
    print("💡 If B < A:")
    print("   → Strategy WORKS! Apply to full document")
    print()
    print("❌ If B ≥ A:")
    print("   → Try more aggressive or different approach")
    print("="*70)
    
    return output_a, output_b

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 5:
        print("Usage: python create_minimal_ab.py <original.docx> <flagged.json> <paraphrased.json> <pages>")
        print("Example: python create_minimal_ab.py testing.docx testing_flagged.json focused_paraphrased.json 14,19,24")
        sys.exit(1)
    
    original = sys.argv[1]
    flagged = sys.argv[2]
    paraphrased = sys.argv[3]
    pages = [int(p.strip()) for p in sys.argv[4].split(',')]
    
    create_minimal_test_files(original, flagged, paraphrased, pages)
