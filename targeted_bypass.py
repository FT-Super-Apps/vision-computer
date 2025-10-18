#!/usr/bin/env python3
"""
Targeted Bypass - ONLY modify flagged texts
Based on manual flag.txt input
"""

import sys
import random
from docx import Document
from difflib import SequenceMatcher

# Homoglyphs mapping (Latin → Cyrillic)
HOMOGLYPHS = {
    'A': 'А', 'B': 'В', 'C': 'С', 'E': 'Е', 'H': 'Н',
    'I': 'І', 'K': 'К', 'M': 'М', 'O': 'О', 'P': 'Р',
    'T': 'Т', 'X': 'Х', 'Y': 'У',
    'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с',
    'x': 'х', 'y': 'у', 'i': 'і', 'k': 'к', 'u': 'υ'
}

def apply_homoglyphs(text, density=0.50):
    """Apply homoglyphs with specified density"""
    if not text:
        return text
    
    result = list(text)
    replaceable = [i for i, c in enumerate(text) if c in HOMOGLYPHS]
    
    if not replaceable:
        return text
    
    num_replace = max(1, int(len(replaceable) * density))
    positions = random.sample(replaceable, min(num_replace, len(replaceable)))
    
    for pos in positions:
        result[pos] = HOMOGLYPHS[result[pos]]
    
    return ''.join(result)

def targeted_bypass(docx_path, flags_file, output_path, density=0.50):
    """
    Find and replace ONLY flagged texts with homoglyphs
    """
    print(f"\n{'='*70}")
    print(f"🎯 TARGETED BYPASS - Manual Flags Only")
    print(f"{'='*70}")
    print(f"Input DOCX : {docx_path}")
    print(f"Flags file : {flags_file}")
    print(f"Output     : {output_path}")
    print(f"Density    : {density*100:.0f}% homoglyphs")
    print(f"{'='*70}\n")
    
    # Load flagged texts
    with open(flags_file, 'r', encoding='utf-8') as f:
        flagged_texts = [line.strip() for line in f if line.strip()]
    
    print(f"📋 Loaded {len(flagged_texts)} flagged texts\n")
    
    # Load document
    doc = Document(docx_path)
    modified_count = 0
    matched_flags = set()
    
    print(f"🔍 Searching and replacing in document...\n")
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        if not para_text:
            continue
        
        # Check against all flagged texts
        for flag_idx, flagged in enumerate(flagged_texts):
            if not flagged:
                continue
            
            # Fuzzy match (handle partial matches)
            similarity = SequenceMatcher(None, para_text.lower(), flagged.lower()).ratio()
            
            # Also check if flagged text is substring
            is_substring = flagged.lower() in para_text.lower()
            
            if similarity > 0.80 or is_substring:
                # Found a match!
                modified_count += 1
                matched_flags.add(flag_idx)
                
                print(f"✅ Match #{modified_count} (similarity: {similarity*100:.1f}%):")
                print(f"   Flagged text: {flagged[:60]}...")
                print(f"   Paragraph #{para_idx}: {para_text[:60]}...")
                
                # Apply homoglyphs to ALL runs in paragraph
                original_full = para.text
                modified_full = apply_homoglyphs(original_full, density=density)
                
                # Clear and rebuild paragraph
                para.clear()
                new_run = para.add_run(modified_full)
                
                # Count changes
                changes = sum(1 for a, b in zip(original_full, modified_full) if a != b)
                print(f"   Changed: {changes}/{len(original_full)} chars ({changes/len(original_full)*100:.1f}%)")
                print(f"   Preview: {modified_full[:60]}...")
                print()
                
                break  # Only match once per paragraph
    
    # Save
    print(f"💾 Saving to: {output_path}")
    doc.save(output_path)
    
    print(f"\n{'='*70}")
    print(f"✅ COMPLETE!")
    print(f"{'='*70}")
    print(f"📊 Statistics:")
    print(f"   Flagged texts total : {len(flagged_texts)}")
    print(f"   Matched & modified  : {len(matched_flags)}")
    print(f"   Paragraphs changed  : {modified_count}")
    print(f"   Density applied     : {density*100:.0f}%")
    print(f"\n📄 Output: {output_path}")
    print(f"{'='*70}\n")
    
    # Show unmatched flags
    unmatched = [i for i in range(len(flagged_texts)) if i not in matched_flags]
    if unmatched:
        print(f"⚠️  Unmatched flags ({len(unmatched)}):")
        for i in unmatched[:5]:
            print(f"   - {flagged_texts[i][:60]}...")
        print()

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python targeted_bypass.py original.docx [flags_file] [output.docx] [density]")
        print("\nExample:")
        print("  python targeted_bypass.py original.docx flag,txt original_targeted.docx 0.50")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    flags_file = sys.argv[2] if len(sys.argv) > 2 else "flag,txt"
    output_file = sys.argv[3] if len(sys.argv) > 3 else "original_targeted.docx"
    density = float(sys.argv[4]) if len(sys.argv) > 4 else 0.50
    
    targeted_bypass(docx_file, flags_file, output_file, density)
